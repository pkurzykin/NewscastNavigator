from __future__ import annotations

import io
import json
import textwrap
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from uuid import NAMESPACE_URL, uuid5

from sqlalchemy import select
from sqlalchemy.orm import Session, aliased

from app.core.config import get_settings
from app.core.version import get_app_version
from app.db.models import Project, ScriptElement, User
from app.schemas.captionpanels_import import (
    CaptionPanelsImportDocument,
    CaptionPanelsImportMeta,
    CaptionPanelsImportSegment,
    CaptionPanelsImportSpeaker,
)
from app.schemas.story_exchange import (
    StoryExchangeDocument,
    StoryExchangeProject,
    StoryExchangeSegment,
    StoryExchangeSegmentFile,
    StoryExchangeSegmentNotes,
    StoryExchangeSource,
    StoryExchangeSpeaker,
)
from app.services.structured_fields import structured_data_from_storage

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas as pdf_canvas
except Exception:
    A4 = None
    TTFont = None
    pdf_canvas = None
    pdfmetrics = None


DOCX_EXPORT_AVAILABLE = Document is not None
PDF_EXPORT_AVAILABLE = A4 is not None and TTFont is not None and pdf_canvas is not None and pdfmetrics is not None

_PDF_FONT_REGISTERED = False
_PDF_FONT_NAME = "Helvetica"

_BLOCK_LABELS = {
    "podvodka": "Подводка",
    "zk": "ЗК",
    "zk_geo": "ЗК+гео",
    "life": "Лайф",
    "snh": "СНХ",
}

_STORY_EXCHANGE_SEMANTIC_TYPES = {
    "podvodka": "voiceover",
    "zk": "voiceover",
    "zk_geo": "voiceover",
    "snh": "sync",
    "life": "sync",
}


class ExportInputNotFoundError(Exception):
    pass


def _format_datetime(value: datetime | None) -> str:
    if value is None:
        return "-"
    return value.strftime("%d.%m.%Y %H:%M")


def _normalize_value(value: str | None) -> str:
    return (value or "").strip()


def _block_label(value: str | None) -> str:
    key = (value or "").strip().lower()
    return _BLOCK_LABELS.get(key, key or "-")


def _normalize_text_lines(raw_value: Any) -> list[str]:
    if isinstance(raw_value, list):
        source = raw_value
    else:
        source = str(raw_value or "").splitlines()
    return [str(item or "").strip() for item in source if str(item or "").strip()]


def build_story_uid(project: Project) -> str:
    return f"story_{project.id}"


def _story_exchange_speaker_uid(story_uid: str, *, name: str, job: str) -> str:
    key = f"{story_uid}:{name.strip().lower()}:{job.strip().lower()}"
    return f"speaker_{uuid5(NAMESPACE_URL, key).hex[:16]}"


def _story_exchange_generated_at() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def _story_exchange_semantic_type(block_type: str | None) -> str:
    normalized_block = (block_type or "").strip().lower()
    return _STORY_EXCHANGE_SEMANTIC_TYPES.get(normalized_block, "voiceover")


def _parse_speaker_lines(raw_value: str | None) -> tuple[str, str]:
    lines = _normalize_text_lines(raw_value)
    if len(lines) >= 2:
        return lines[0], lines[1]
    if len(lines) == 1:
        return lines[0], ""
    return "", ""


def _resolve_export_root() -> Path:
    root = Path(get_settings().export_root).expanduser()
    if not root.is_absolute():
        root = (Path.cwd() / root).resolve()
    root.mkdir(parents=True, exist_ok=True)
    return root


def persist_export_bytes(
    *,
    project_id: int,
    file_name: str,
    content: bytes,
) -> Path:
    timestamp = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
    export_root = _resolve_export_root()
    project_dir = export_root / "projects" / str(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)

    source_path = Path(file_name)
    target_path = project_dir / f"{source_path.stem}-{timestamp}{source_path.suffix}"
    target_path.write_bytes(content)
    return target_path


def build_story_exchange_payload(db: Session, project_id: int) -> dict[str, Any]:
    project = db.execute(select(Project).where(Project.id == project_id)).scalar_one_or_none()
    if project is None:
        raise ExportInputNotFoundError("Проект не найден")
    story_uid = build_story_uid(project)
    elements = db.execute(
        select(ScriptElement)
        .where(ScriptElement.project_id == project_id)
        .order_by(ScriptElement.order_index.asc(), ScriptElement.id.asc())
    ).scalars().all()

    speakers_by_id: dict[str, StoryExchangeSpeaker] = {}
    segments: list[StoryExchangeSegment] = []

    for index, item in enumerate(elements, start=1):
        block_type = (item.block_type or "zk").strip().lower()
        structured_data = structured_data_from_storage(
            block_type=block_type,
            text=item.text,
            content_json=item.content_json,
        )
        text_lines = _normalize_text_lines(item.text)
        geo: str | None = None

        if block_type == "zk_geo":
            geo = _normalize_value(str(structured_data.get("geo") or "")) or None
            text_lines = _normalize_text_lines(structured_data.get("text_lines"))

        speaker_id: str | None = None
        if block_type == "snh":
            speaker_name, speaker_job = _parse_speaker_lines(item.speaker_text)
            if speaker_name or speaker_job:
                speaker_id = _story_exchange_speaker_uid(
                    story_uid,
                    name=speaker_name,
                    job=speaker_job,
                )
                if speaker_id not in speakers_by_id:
                    speakers_by_id[speaker_id] = StoryExchangeSpeaker(
                        speaker_id=speaker_id,
                        name=speaker_name,
                        job=speaker_job,
                    )

        segments.append(
            StoryExchangeSegment(
                segment_uid=item.segment_uid,
                order=index,
                block_type=block_type,
                semantic_type=_story_exchange_semantic_type(block_type),
                text="\n".join(text_lines),
                text_lines=text_lines,
                geo=geo,
                speaker_id=speaker_id,
                file=StoryExchangeSegmentFile(
                    name=_normalize_value(item.file_name),
                    tc_in=_normalize_value(item.tc_in),
                    tc_out=_normalize_value(item.tc_out),
                ),
                notes=StoryExchangeSegmentNotes(
                    on_screen=_normalize_value(item.additional_comment),
                ),
            )
        )

    document = StoryExchangeDocument(
        schema_version=1,
        story_uid=story_uid,
        generated_at=_story_exchange_generated_at(),
        source=StoryExchangeSource(
            system="newscastnavigator",
            version=get_app_version(),
        ),
        project=StoryExchangeProject(
            id=project.id,
            title=project.title,
            rubric=_normalize_value(project.rubric),
            planned_duration=_normalize_value(project.planned_duration),
            status=_normalize_value(project.status),
        ),
        speakers=list(speakers_by_id.values()),
        segments=segments,
    )
    return document.model_dump(mode="json", by_alias=True, exclude_none=True)


def generate_story_exchange_bytes(payload: dict[str, Any]) -> bytes:
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def build_captionpanels_import_payload(db: Session, project_id: int) -> dict[str, Any]:
    story_payload = build_story_exchange_payload(db, project_id)
    project_payload = story_payload["project"]
    story_segments = story_payload["segments"]
    story_speakers = story_payload.get("speakers", [])

    speakers = [
        CaptionPanelsImportSpeaker(
            id=str(item["speakerId"]),
            name=str(item.get("name") or ""),
            job=str(item.get("job") or ""),
        )
        for item in story_speakers
    ]

    segments: list[CaptionPanelsImportSegment] = []
    for item in story_segments:
        segment_uid = str(item["segmentUid"])
        semantic_type = str(item.get("semanticType") or "")
        block_type = str(item.get("blockType") or "")
        geo_text = str(item.get("geo") or "").strip()

        if block_type == "zk_geo" and geo_text:
            segments.append(
                CaptionPanelsImportSegment(
                    id=f"{segment_uid}:geo",
                    type="geotag",
                    text=geo_text,
                )
            )

        target_type = "synch" if semantic_type == "sync" else "voiceover"
        segments.append(
            CaptionPanelsImportSegment(
                id=segment_uid,
                type=target_type,
                text=str(item.get("text") or ""),
                speaker_id=item.get("speakerId"),
            )
        )

    document = CaptionPanelsImportDocument(
        meta=CaptionPanelsImportMeta(
            title=str(project_payload.get("title") or ""),
            rubric=str(project_payload.get("rubric") or ""),
        ),
        speakers=speakers,
        segments=segments,
    )
    return document.model_dump(mode="json", by_alias=True, exclude_none=True)


def generate_captionpanels_import_bytes(payload: dict[str, Any]) -> bytes:
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def fetch_export_payload(db: Session, project_id: int) -> dict[str, Any]:
    author_user = aliased(User)
    executor_user = aliased(User)
    proofreader_user = aliased(User)
    row = db.execute(
        select(Project, author_user.username, executor_user.username, proofreader_user.username)
        .outerjoin(author_user, author_user.id == Project.author_user_id)
        .outerjoin(executor_user, executor_user.id == Project.executor_user_id)
        .outerjoin(proofreader_user, proofreader_user.id == Project.proofreader_user_id)
        .where(Project.id == project_id)
    ).first()
    if row is None:
        raise ExportInputNotFoundError("Проект не найден")

    project: Project = row[0]
    author_username: str | None = row[1]
    executor_username: str | None = row[2]
    proofreader_username: str | None = row[3]
    elements = db.execute(
        select(ScriptElement)
        .where(ScriptElement.project_id == project_id)
        .order_by(ScriptElement.order_index.asc(), ScriptElement.id.asc())
    ).scalars().all()

    payload_rows = []
    for index, item in enumerate(elements, start=1):
        structured_data = structured_data_from_storage(
            block_type=item.block_type,
            text=item.text,
            content_json=item.content_json,
        )
        export_text = _normalize_value(item.text)
        if (item.block_type or "").strip().lower() == "zk_geo":
            geo = _normalize_value(str(structured_data.get("geo") or ""))
            text_lines = structured_data.get("text_lines") or []
            text_value = "\n".join(str(line) for line in text_lines if str(line).strip())
            export_text = "\n".join(
                part for part in [f"Гео: {geo}" if geo else "", text_value] if part
            )
        payload_rows.append(
            {
                "index": index,
                "block": _block_label(item.block_type),
                "text": export_text,
                "speaker_text": _normalize_value(item.speaker_text),
                "file_name": _normalize_value(item.file_name),
                "tc_in": _normalize_value(item.tc_in),
                "tc_out": _normalize_value(item.tc_out),
                "additional_comment": _normalize_value(item.additional_comment),
            }
        )

    return {
        "project_id": project.id,
        "title": project.title,
        "status": project.status,
        "rubric": _normalize_value(project.rubric),
        "planned_duration": _normalize_value(project.planned_duration),
        "project_note": _normalize_value(project.project_note),
        "author_username": author_username or "-",
        "executor_username": executor_username or "-",
        "proofreader_username": proofreader_username or "-",
        "created_at": _format_datetime(project.created_at),
        "rows": payload_rows,
    }


def generate_docx_bytes(payload: dict[str, Any]) -> bytes:
    if not DOCX_EXPORT_AVAILABLE:
        raise RuntimeError("Для DOCX установите пакет python-docx")

    document = Document()
    document.add_heading("Newscast Navigator", level=1)
    document.add_paragraph(f"Проект: {payload['title']}")
    document.add_paragraph(
        f"Рубрика: {payload['rubric'] or '-'} | Хронометраж: {payload['planned_duration'] or '-'}"
    )
    document.add_paragraph(
        f"Автор: {payload['author_username']} | Создан: {payload['created_at']}"
    )
    document.add_paragraph(
        f"Исполнитель: {payload['executor_username']} | Корректор: {payload['proofreader_username']}"
    )
    if payload["project_note"]:
        document.add_paragraph(f"Комментарий проекта: {payload['project_note']}")

    headers = ["№", "Блок", "Текст", "Титр", "Имя файла", "TC IN", "TC OUT", "Комментарий"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for col, header in enumerate(headers):
        table.rows[0].cells[col].text = header

    for row in payload["rows"]:
        cells = table.add_row().cells
        cells[0].text = str(row["index"])
        cells[1].text = row["block"]
        cells[2].text = row["text"]
        cells[3].text = row["speaker_text"]
        cells[4].text = row["file_name"]
        cells[5].text = row["tc_in"]
        cells[6].text = row["tc_out"]
        cells[7].text = row["additional_comment"]

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _ensure_pdf_font() -> str:
    global _PDF_FONT_REGISTERED, _PDF_FONT_NAME

    if not PDF_EXPORT_AVAILABLE:
        return "Helvetica"
    if _PDF_FONT_REGISTERED:
        return _PDF_FONT_NAME

    candidates = [
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        Path("/Library/Fonts/Arial Unicode.ttf"),
    ]
    for candidate in candidates:
        if not candidate.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont("NewscastUnicode", str(candidate)))
            _PDF_FONT_NAME = "NewscastUnicode"
            _PDF_FONT_REGISTERED = True
            return _PDF_FONT_NAME
        except Exception:
            continue

    _PDF_FONT_REGISTERED = True
    _PDF_FONT_NAME = "Helvetica"
    return _PDF_FONT_NAME


def generate_pdf_bytes(payload: dict[str, Any]) -> bytes:
    if not PDF_EXPORT_AVAILABLE:
        raise RuntimeError("Для PDF установите пакет reportlab")

    buffer = io.BytesIO()
    pdf = pdf_canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin_x = 40
    margin_top = 40
    y = height - margin_top
    line_height = 14
    font_name = _ensure_pdf_font()
    pdf.setFont(font_name, 10)

    lines = [
        "Newscast Navigator",
        f"Проект: {payload['title']}",
        f"Рубрика: {payload['rubric'] or '-'} | Хронометраж: {payload['planned_duration'] or '-'}",
        f"Автор: {payload['author_username']} | Создан: {payload['created_at']}",
        f"Исполнитель: {payload['executor_username']} | Корректор: {payload['proofreader_username']}",
        f"Комментарий проекта: {payload['project_note'] or '-'}",
        "",
    ]
    for row in payload["rows"]:
        lines.extend(
            [
                f"#{row['index']} | {row['block']}",
                f"Текст: {row['text'] or '-'}",
                f"Титр: {row['speaker_text'] or '-'}",
                f"Файл: {row['file_name'] or '-'} | TC: {row['tc_in'] or '-'} - {row['tc_out'] or '-'}",
                f"Комментарий: {row['additional_comment'] or '-'}",
                "",
            ]
        )

    max_width_chars = 110
    for line in lines:
        wrapped_lines = textwrap.wrap(line, width=max_width_chars) if line else [""]
        for wrapped in wrapped_lines:
            if y <= 35:
                pdf.showPage()
                pdf.setFont(font_name, 10)
                y = height - margin_top
            pdf.drawString(margin_x, y, wrapped)
            y -= line_height

    pdf.save()
    return buffer.getvalue()
