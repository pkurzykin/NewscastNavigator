from __future__ import annotations

from collections.abc import Mapping, Sequence
from dataclasses import dataclass, replace
from io import BytesIO
import re
import unicodedata
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import _Cell, _Row, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.services.scenario_docx_snapshot import ScenarioDocxRow, ScenarioDocxSnapshot


DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

ALLOWED_FONTS = {"PT Sans", "Arial", "Georgia", "Times New Roman", "Roboto Slab"}
ALLOWED_FILLS = {"#ffffff", "#ffff00", "#ff0000", "#00ff00", "#0000ff", "#ffa500"}

_COLUMN_WIDTHS = (Cm(9.68), Cm(5.26), Cm(2.55))
_WINDOWS_RESERVED = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    *(f"COM{number}" for number in range(1, 10)),
    *(f"LPT{number}" for number in range(1, 10)),
}
_INVALID_FILENAME_CHARACTERS = frozenset('/\\:*?"<>|')


@dataclass(frozen=True)
class DocxRunStyle:
    font_family: str
    bold: bool
    italic: bool
    strikethrough: bool
    fill_color: str | None


@dataclass(frozen=True)
class DocxTextRun:
    text: str
    style: DocxRunStyle
    hard_break: bool = False


@dataclass(frozen=True)
class DocxParagraph:
    runs: tuple[DocxTextRun, ...]


def _default_target_style(block_type: str, target: str) -> DocxRunStyle:
    return DocxRunStyle(
        font_family="PT Sans",
        bold=block_type == "snh" and target != "text",
        italic=(
            block_type == "life"
            or (block_type == "zk_geo" and target == "geo")
            or block_type == "snh"
        ),
        strikethrough=False,
        fill_color=None,
    )


def _as_mapping(value: object) -> Mapping[str, Any]:
    return value if isinstance(value, Mapping) else {}


def _as_nodes(value: object) -> Sequence[object]:
    if isinstance(value, Sequence) and not isinstance(value, (str, bytes, bytearray)):
        return value
    return ()


def _target_style(row: ScenarioDocxRow, target: str) -> DocxRunStyle:
    style = _default_target_style(row.block_type, target)
    target_formatting = _as_mapping(_as_mapping(row.formatting.get("targets")).get(target))

    font_family = target_formatting.get("font_family")
    if isinstance(font_family, str) and font_family in ALLOWED_FONTS:
        style = replace(style, font_family=font_family)
    for field_name in ("bold", "italic", "strikethrough"):
        value = target_formatting.get(field_name)
        if isinstance(value, bool):
            style = replace(style, **{field_name: value})
    fill_color = target_formatting.get("fill_color")
    if isinstance(fill_color, str) and fill_color.casefold() in ALLOWED_FILLS:
        normalized_fill = fill_color.casefold()
        style = replace(
            style,
            fill_color=None if normalized_fill == "#ffffff" else normalized_fill,
        )
    return style


def _style_with_marks(style: DocxRunStyle, raw_marks: object) -> DocxRunStyle:
    result = style
    for raw_mark in _as_nodes(raw_marks):
        mark = _as_mapping(raw_mark)
        mark_type = mark.get("type")
        if mark_type == "bold":
            result = replace(result, bold=True)
        elif mark_type == "italic":
            result = replace(result, italic=True)
        elif mark_type == "strike":
            result = replace(result, strikethrough=True)
        elif mark_type == "textStyle":
            font_family = _as_mapping(mark.get("attrs")).get("fontFamily")
            if isinstance(font_family, str) and font_family in ALLOWED_FONTS:
                result = replace(result, font_family=font_family)
        elif mark_type == "highlight":
            color = _as_mapping(mark.get("attrs")).get("color")
            if isinstance(color, str) and color.casefold() in ALLOWED_FILLS:
                normalized_color = color.casefold()
                result = replace(
                    result,
                    fill_color=None if normalized_color == "#ffffff" else normalized_color,
                )
    return result


def _parse_tiptap_doc(
    raw_doc: object,
    default_style: DocxRunStyle,
) -> tuple[DocxParagraph, ...] | None:
    doc = _as_mapping(raw_doc)
    if doc.get("type") != "doc" or not isinstance(doc.get("content"), Sequence):
        return None

    paragraphs: list[list[DocxTextRun]] = []

    def walk(node_value: object, current: list[DocxTextRun] | None = None) -> None:
        node = _as_mapping(node_value)
        node_type = node.get("type")
        if node_type == "paragraph":
            paragraph: list[DocxTextRun] = []
            paragraphs.append(paragraph)
            for child in _as_nodes(node.get("content")):
                walk(child, paragraph)
            return
        if node_type == "hardBreak":
            paragraph = current
            if paragraph is None:
                paragraph = []
                paragraphs.append(paragraph)
            paragraph.append(DocxTextRun("", default_style, hard_break=True))
            return
        if node_type == "text":
            text = node.get("text")
            if not isinstance(text, str) or not text:
                return
            paragraph = current
            if paragraph is None:
                paragraph = []
                paragraphs.append(paragraph)
            paragraph.append(
                DocxTextRun(text, _style_with_marks(default_style, node.get("marks")))
            )
            return
        for child in _as_nodes(node.get("content")):
            walk(child, current)

    for child in _as_nodes(doc.get("content")):
        walk(child)
    return tuple(DocxParagraph(tuple(runs)) for runs in paragraphs)


def _visible_text(paragraphs: tuple[DocxParagraph, ...]) -> str:
    return "\n".join(
        "".join("\n" if run.hard_break else run.text for run in paragraph.runs)
        for paragraph in paragraphs
    )


def _plain_paragraphs(text: str, style: DocxRunStyle) -> tuple[DocxParagraph, ...]:
    if not text:
        return ()
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    return tuple(
        DocxParagraph((DocxTextRun(line, style),))
        for line in normalized.split("\n")
    )


def _target_text(row: ScenarioDocxRow, target: str) -> str:
    if target == "text":
        return row.text
    if target == "geo":
        value = row.structured_data.get("geo", "")
        return value if isinstance(value, str) else ""
    fio, position = (row.speaker_text.split("\n", 1) + [""])[:2]
    return fio if target == "speaker_fio" else position


def _target_paragraphs(row: ScenarioDocxRow, target: str) -> tuple[DocxParagraph, ...]:
    canonical_text = _target_text(row, target)
    if not canonical_text:
        return ()
    style = _target_style(row, target)
    rich_target = _as_mapping(_as_mapping(row.rich_text.get("targets")).get(target))
    parsed = _parse_tiptap_doc(rich_target.get("doc"), style)
    if parsed is not None and _visible_text(parsed) == canonical_text:
        return parsed
    return _plain_paragraphs(canonical_text, style)


def _set_cell_shading(cell: _Cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)


def _set_table_borders(table: Table) -> None:
    properties = table._tbl.tblPr
    existing = properties.find(qn("w:tblBorders"))
    if existing is not None:
        properties.remove(existing)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        borders.append(border)
    properties.append(borders)


def _set_cell_width(cell: _Cell, width: int) -> None:
    cell.width = width
    properties = cell._tc.get_or_add_tcPr()
    tc_width = properties.find(qn("w:tcW"))
    if tc_width is None:
        tc_width = OxmlElement("w:tcW")
        properties.append(tc_width)
    tc_width.set(qn("w:type"), "dxa")
    tc_width.set(qn("w:w"), str(width.twips))


def _set_repeat_table_header(row: _Row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:tblHeader")) is None:
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        properties.append(repeat)


def _set_run_fill(run: Run, color: str | None) -> None:
    properties = run._r.get_or_add_rPr()
    existing = properties.find(qn("w:shd"))
    if existing is not None:
        properties.remove(existing)
    if color is None:
        return
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), color.removeprefix("#").upper())
    properties.append(shading)


def _xml_safe_text(value: str) -> str:
    """Replace each XML 1.0-invalid code point with U+FFFD."""

    def is_valid(character: str) -> bool:
        code_point = ord(character)
        return (
            code_point in {0x09, 0x0A, 0x0D}
            or 0x20 <= code_point <= 0xD7FF
            or 0xE000 <= code_point <= 0xFFFD
            or 0x10000 <= code_point <= 0x10FFFF
        )

    return "".join(
        character if is_valid(character) else "\uFFFD"
        for character in value
    )


def _add_safe_run(paragraph: Paragraph, text: str = "") -> Run:
    return paragraph.add_run(_xml_safe_text(text))


def _append_safe_text(run: Run, text: str) -> None:
    run.add_text(_xml_safe_text(text))


def _apply_run_style(run: Run, style: DocxRunStyle) -> None:
    run.font.name = style.font_family
    run.font.size = Pt(12)
    run.bold = style.bold
    run.italic = style.italic
    run.font.strike = style.strikethrough
    fonts = run._r.get_or_add_rPr().get_or_add_rFonts()
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{slot}"), style.font_family)
    _set_run_fill(run, style.fill_color)


class _CellWriter:
    def __init__(
        self,
        cell: _Cell,
        *,
        alignment: WD_PARAGRAPH_ALIGNMENT | None = None,
    ) -> None:
        cell.text = ""
        cell.paragraphs[0].clear()
        self._cell = cell
        self._first = True
        self._alignment = alignment

    def paragraph(self) -> Paragraph:
        if self._first:
            self._first = False
            paragraph = self._cell.paragraphs[0]
        else:
            paragraph = self._cell.add_paragraph()
        paragraph.alignment = self._alignment
        return paragraph

    def append(self, paragraphs: tuple[DocxParagraph, ...]) -> None:
        for source in paragraphs:
            paragraph = self.paragraph()
            for source_run in source.runs:
                run = _add_safe_run(paragraph, source_run.text)
                if source_run.hard_break:
                    run.add_break()
                _apply_run_style(run, source_run.style)

    def append_plain(self, text: str, style: DocxRunStyle) -> None:
        if text:
            self.append((DocxParagraph((DocxTextRun(text, style),)),))


def _write_simple_text(
    cell: _Cell,
    text: str,
    *,
    bold: bool = False,
    alignment: WD_PARAGRAPH_ALIGNMENT | None = None,
) -> None:
    writer = _CellWriter(cell, alignment=alignment)
    writer.append_plain(
        text,
        replace(_default_target_style("", "text"), bold=bold),
    )


def _write_header_texts(cell: _Cell, *texts: str) -> None:
    writer = _CellWriter(cell, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    style = replace(_default_target_style("", "text"), bold=True)
    for text in texts:
        writer.append_plain(text, style)


def _write_body_row(table_row: _Row, source: ScenarioDocxRow) -> None:
    text_writer = _CellWriter(
        table_row.cells[0],
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    video_writer = _CellWriter(table_row.cells[1])
    sound_writer = _CellWriter(table_row.cells[2])

    if source.block_type in {"podvodka", "zk"}:
        text_writer.append(_target_paragraphs(source, "text"))
    elif source.block_type == "zk_geo":
        text_writer.append(_target_paragraphs(source, "geo"))
        text_writer.append(_target_paragraphs(source, "text"))
    elif source.block_type == "snh":
        text_writer.append(_target_paragraphs(source, "speaker_fio"))
        text_writer.append(_target_paragraphs(source, "speaker_position"))
        text_writer.append(_target_paragraphs(source, "text"))
    elif source.block_type == "life":
        sound_writer.append(_target_paragraphs(source, "text"))
    else:
        raise ValueError(f"Unsupported scenario block type: {source.block_type}")

    plain_style = _default_target_style("", "text")
    for bundle in source.file_bundles:
        tc = (
            f"{bundle.tc_in}–{bundle.tc_out}"
            if bundle.tc_in and bundle.tc_out
            else bundle.tc_in or bundle.tc_out
        )
        if not bundle.file_name and not tc:
            continue
        paragraph = video_writer.paragraph()
        if bundle.file_name:
            run = _add_safe_run(paragraph, bundle.file_name)
            _apply_run_style(run, plain_style)
        if tc:
            run = _add_safe_run(paragraph)
            if bundle.file_name:
                run.add_break()
            _append_safe_text(run, tc)
            _apply_run_style(run, plain_style)
    video_writer.append_plain(source.additional_comment, plain_style)


def render_scenario_docx(snapshot: ScenarioDocxSnapshot) -> BytesIO:
    document = Document()
    core = document.core_properties
    core.author = ""
    core.last_modified_by = ""
    core.comments = ""
    core.keywords = ""
    core.subject = ""
    core.title = ""
    last_modified_by = core._element.find(qn("cp:lastModifiedBy"))
    if last_modified_by is not None:
        core._element.remove(last_modified_by)

    normal = document.styles["Normal"]
    normal.font.name = "PT Sans"
    normal.font.size = Pt(12)
    normal_fonts = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal_fonts.set(qn(f"w:{slot}"), "PT Sans")

    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.right_margin = Cm(1.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)

    table = document.add_table(rows=3 + len(snapshot.rows), cols=3)
    table.autofit = False
    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table._tbl.tblPr.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(Cm(17.49).twips))
    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
    for index, width in enumerate(_COLUMN_WIDTHS):
        grid_columns[index].set(qn("w:w"), str(width.twips))
        for row in table.rows:
            _set_cell_width(row.cells[index], width)

    title_rubric = table.rows[0].cells[0].merge(table.rows[0].cells[2])
    _write_header_texts(title_rubric, snapshot.title, snapshot.rubric_name)
    _set_cell_shading(title_rubric, "B6DDE8")

    duration_text = (
        f"Хронометраж {snapshot.duration_text}"
        if snapshot.duration_text
        else "Хронометраж —"
    )
    duration = table.rows[1].cells[0].merge(table.rows[1].cells[2])
    _write_header_texts(duration, duration_text)
    _set_cell_shading(duration, "B6DDE8")

    header = table.rows[2]
    for cell, text in zip(
        header.cells,
        ("Текст в титре", "В кадре", "Звук"),
        strict=True,
    ):
        _write_simple_text(
            cell,
            text,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_shading(cell, "B6DDE8")
    for row in table.rows[:3]:
        _set_repeat_table_header(row)

    for table_row, source in zip(table.rows[3:], snapshot.rows, strict=True):
        _write_body_row(table_row, source)

    _set_table_borders(table)
    for row in table.rows[:3]:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in table.rows[3:]:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


def safe_docx_filename(title: str, story_id: int) -> tuple[str, str]:
    fallback = f"Scenario-{story_id}.docx"
    cleaned_characters = [
        "-"
        if character in _INVALID_FILENAME_CHARACTERS
        or unicodedata.category(character).startswith("C")
        else character
        for character in str(title)
    ]
    stem = "".join(cleaned_characters)
    stem = re.sub(r"\.{2,}", "-", stem)
    stem = re.sub(r"\s+", " ", stem)
    stem = re.sub(r"-+", "-", stem)
    stem = re.sub(r"\s*-\s*", "-", stem)
    stem = stem.strip(" .-")[:120].rstrip(" .")
    reserved_base = stem.split(".", 1)[0].upper()
    if not stem or reserved_base in _WINDOWS_RESERVED:
        stem = f"Сценарий-{story_id}"
    return fallback, f"{stem}.docx"
