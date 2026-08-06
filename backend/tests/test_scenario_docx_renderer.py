from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path
import tempfile
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
import pytest

from app.services.scenario_docx_renderer import (
    DOCX_CONTENT_TYPE,
    render_scenario_docx,
    safe_docx_filename,
)
from app.services.scenario_docx_snapshot import (
    DocxFileBundle,
    ScenarioDocxRow,
    ScenarioDocxSnapshot,
)


def _row(
    block_type: str,
    text: str,
    *,
    speaker_text: str = "",
    additional_comment: str = "",
    structured_data: dict[str, object] | None = None,
    formatting: dict[str, object] | None = None,
    rich_text: dict[str, object] | None = None,
    file_bundles: tuple[DocxFileBundle, ...] = (),
) -> ScenarioDocxRow:
    return ScenarioDocxRow(
        block_type=block_type,
        text=text,
        speaker_text=speaker_text,
        additional_comment=additional_comment,
        structured_data=structured_data or {},
        formatting=formatting or {},
        rich_text=rich_text or {},
        file_bundles=file_bundles,
    )


def _snapshot(*rows: ScenarioDocxRow) -> ScenarioDocxSnapshot:
    return ScenarioDocxSnapshot(
        story_id=73,
        title="Синтетический выпуск",
        rubric_id=9,
        rubric_name="Учебная рубрика",
        duration_text="12:34",
        revision=4,
        rows=tuple(rows),
    )


def _fixture_snapshot() -> ScenarioDocxSnapshot:
    return _snapshot(
        _row(
            "podvodka",
            "Жирный курсив\nВложенный\nВторой абзац",
            rich_text={
                "targets": {
                    "text": {
                        "doc": {
                            "type": "doc",
                            "content": [
                                {
                                    "type": "paragraph",
                                    "content": [
                                        {
                                            "type": "text",
                                            "text": "Жирный",
                                            "marks": [{"type": "bold"}],
                                        },
                                        {
                                            "type": "text",
                                            "text": " курсив",
                                            "marks": [
                                                {"type": "italic"},
                                                {"type": "strike"},
                                                {
                                                    "type": "textStyle",
                                                    "attrs": {
                                                        "fontFamily": "Arial",
                                                        "style": "font-size:999px",
                                                    },
                                                },
                                                {
                                                    "type": "highlight",
                                                    "attrs": {
                                                        "color": "#ffff00",
                                                        "style": "position:fixed",
                                                    },
                                                },
                                            ],
                                        },
                                        {"type": "hardBreak"},
                                        {
                                            "type": "unsupported-wrapper",
                                            "attrs": {"style": "color:red"},
                                            "content": [
                                                {
                                                    "type": "text",
                                                    "text": "Вложенный",
                                                    "marks": [
                                                        {
                                                            "type": "textStyle",
                                                            "attrs": {
                                                                "fontFamily": "url(file:///private/font.ttf)"
                                                            },
                                                        },
                                                        {
                                                            "type": "highlight",
                                                            "attrs": {
                                                                "color": "expression(alert(1))"
                                                            },
                                                        },
                                                    ],
                                                }
                                            ],
                                        },
                                    ],
                                },
                                {
                                    "type": "paragraph",
                                    "content": [
                                        {
                                            "type": "text",
                                            "text": "Второй абзац",
                                            "attrs": {"html": "<w:tbl/>"},
                                        }
                                    ],
                                },
                            ],
                        }
                    }
                }
            },
            file_bundles=(
                DocxFileBundle("first-reference.mov", "00:01", "00:03"),
                DocxFileBundle("second-reference.mov", "00:05", ""),
            ),
            additional_comment="Отдельный синтетический комментарий",
        ),
        _row(
            "zk",
            "Текст за кадром",
            formatting={
                "targets": {
                    "text": {
                        "font_family": "Georgia",
                        "bold": True,
                        "italic": False,
                        "strikethrough": True,
                        "fill_color": "#ff0000",
                    }
                }
            },
            file_bundles=(DocxFileBundle("only-timecode.mov", "", "00:08"),),
        ),
        _row(
            "zk_geo",
            "Текст географического блока",
            structured_data={"geo": "Тестоград"},
        ),
        _row("snh", "Речь синтетического героя", speaker_text="Тестов Тест\nЭксперт"),
        _row("life", "Натуральный звук"),
    )


def _nonempty_paragraphs(cell: object) -> list[object]:
    return [paragraph for paragraph in cell.paragraphs if paragraph.text]


def _shading_fill(element: object) -> str | None:
    shading = element.find(qn("w:shd"))
    return None if shading is None else shading.get(qn("w:fill"))


def _tree_listing(path: Path) -> tuple[str, ...]:
    if not path.exists():
        return ()
    return tuple(sorted(str(item.relative_to(path)) for item in path.rglob("*")))


def test_renderer_builds_a4_table_layout_and_all_five_block_mappings() -> None:
    buffer = render_scenario_docx(_fixture_snapshot())

    assert DOCX_CONTENT_TYPE == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert isinstance(buffer, BytesIO)
    assert buffer.tell() == 0
    document = Document(buffer)
    assert len(document.sections) == 1
    section = document.sections[0]
    assert section.orientation == WD_ORIENT.PORTRAIT
    assert section.page_width.cm == pytest.approx(21.0, abs=0.01)
    assert section.page_height.cm == pytest.approx(29.7, abs=0.01)
    assert section.top_margin.cm == pytest.approx(2.0, abs=0.01)
    assert section.right_margin.cm == pytest.approx(1.5, abs=0.01)
    assert section.bottom_margin.cm == pytest.approx(2.0, abs=0.01)
    assert section.left_margin.cm == pytest.approx(3.0, abs=0.01)

    assert len(document.tables) == 1
    table = document.tables[0]
    assert table.autofit is False
    assert len(table.rows) == 9
    grid_widths = [
        int(column.get(qn("w:w")))
        for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
    ]
    assert grid_widths == pytest.approx([5488, 2982, 1446], abs=1)
    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    assert table_width is not None
    assert int(table_width.get(qn("w:w"))) == pytest.approx(9916, abs=1)

    for row in table.rows[:3]:
        assert row.cells[0]._tc is row.cells[1]._tc is row.cells[2]._tc
    assert table.rows[0].cells[0].text == "Синтетический выпуск"
    assert table.rows[1].cells[0].text == "Учебная рубрика"
    assert table.rows[2].cells[0].text == "Хронометраж 12:34"
    assert [
        _shading_fill(row.cells[0]._tc.get_or_add_tcPr())
        for row in table.rows[:3]
    ] == ["B6DDE8", "B6DDE8", "B6DDE8"]

    header = table.rows[3]
    assert [cell.text for cell in header.cells] == [
        "Текст в титре",
        "В кадре",
        "Звук",
    ]
    assert [_shading_fill(cell._tc.get_or_add_tcPr()) for cell in header.cells] == [
        "B6DDE8",
        "B6DDE8",
        "B6DDE8",
    ]
    assert header._tr.get_or_add_trPr().find(qn("w:tblHeader")) is not None

    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None
    assert {
        border.tag.rsplit("}", 1)[-1]: border.get(qn("w:sz"))
        for border in borders
    } == {
        "top": "4",
        "left": "4",
        "bottom": "4",
        "right": "4",
        "insideH": "4",
        "insideV": "4",
    }
    assert all(
        cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.TOP
        for row in table.rows
        for cell in row.cells
    )
    assert all(row._tr.find("./w:trPr/w:trHeight", row._tr.nsmap) is None for row in table.rows)

    body = table.rows[4:]
    assert len(body) == 5
    assert [paragraph.text for paragraph in _nonempty_paragraphs(body[0].cells[0])] == [
        "Жирный курсив\nВложенный",
        "Второй абзац",
    ]
    assert [paragraph.text for paragraph in _nonempty_paragraphs(body[0].cells[1])] == [
        "first-reference.mov\n00:01–00:03",
        "second-reference.mov\n00:05",
        "Отдельный синтетический комментарий",
    ]
    assert [paragraph.text for paragraph in _nonempty_paragraphs(body[1].cells[1])] == [
        "only-timecode.mov\n00:08"
    ]
    assert [paragraph.text for paragraph in _nonempty_paragraphs(body[2].cells[0])] == [
        "Тестоград",
        "Текст географического блока",
    ]
    assert [paragraph.text for paragraph in _nonempty_paragraphs(body[3].cells[0])] == [
        "Тестов Тест",
        "Эксперт",
        "Речь синтетического героя",
    ]
    assert body[4].cells[0].text == ""
    assert body[4].cells[2].text == "Натуральный звук"


def test_renderer_rejects_unknown_block_type() -> None:
    snapshot = _snapshot(_row("corrupted", "Этот текст нельзя экспортировать"))

    with pytest.raises(ValueError, match="Unsupported scenario block type: corrupted"):
        render_scenario_docx(snapshot)


def test_renderer_preserves_whitelisted_styles_and_applies_safe_defaults() -> None:
    document = Document(render_scenario_docx(_fixture_snapshot()))
    assert document.styles["Normal"].font.name == "PT Sans"
    assert document.styles["Normal"].font.size == Pt(12)
    body = document.tables[0].rows[4:]

    rich_paragraphs = _nonempty_paragraphs(body[0].cells[0])
    assert len(rich_paragraphs) == 2
    assert [run.text for run in rich_paragraphs[0].runs] == [
        "Жирный",
        " курсив",
        "\n",
        "Вложенный",
    ]
    bold_run, marked_run, _break_run, unsafe_run = rich_paragraphs[0].runs
    assert bold_run.text == "Жирный"
    assert bold_run.bold is True
    assert marked_run.text == " курсив"
    assert marked_run.italic is True
    assert marked_run.font.strike is True
    assert marked_run.font.name == "Arial"
    fonts = marked_run._r.get_or_add_rPr().find(qn("w:rFonts"))
    assert fonts is not None
    assert {fonts.get(qn(f"w:{slot}")) for slot in ("ascii", "hAnsi", "eastAsia", "cs")} == {
        "Arial"
    }
    assert _shading_fill(marked_run._r.get_or_add_rPr()) == "FFFF00"

    assert unsafe_run.text == "Вложенный"
    assert unsafe_run.font.name == "PT Sans"
    assert _shading_fill(unsafe_run._r.get_or_add_rPr()) is None
    assert all(run.font.size == Pt(12) for paragraph in rich_paragraphs for run in paragraph.runs)
    document_xml = document.part.blob.decode("utf-8")
    assert "font-size:999px" not in document_xml
    assert "position:fixed" not in document_xml
    assert "url(file:///private/font.ttf)" not in document_xml
    assert "expression(alert(1))" not in document_xml
    assert "&lt;w:tbl/&gt;" not in document_xml

    formatted_run = _nonempty_paragraphs(body[1].cells[0])[0].runs[0]
    assert formatted_run.font.name == "Georgia"
    assert formatted_run.bold is True
    assert formatted_run.italic is False
    assert formatted_run.font.strike is True
    assert _shading_fill(formatted_run._r.get_or_add_rPr()) == "FF0000"

    geo_run = _nonempty_paragraphs(body[2].cells[0])[0].runs[0]
    geo_text_run = _nonempty_paragraphs(body[2].cells[0])[1].runs[0]
    assert geo_run.italic is True
    assert geo_text_run.italic is False
    speaker_fio, speaker_position, speaker_text = _nonempty_paragraphs(body[3].cells[0])
    assert speaker_fio.runs[0].bold is True and speaker_fio.runs[0].italic is True
    assert speaker_position.runs[0].bold is True and speaker_position.runs[0].italic is True
    assert speaker_text.runs[0].bold is False and speaker_text.runs[0].italic is True
    assert _nonempty_paragraphs(body[4].cells[2])[0].runs[0].italic is True


def test_rich_marks_override_persisted_target_formatting_and_white_removes_shading() -> None:
    snapshot = _snapshot(
        _row(
            "life",
            "Белая метка и база",
            formatting={
                "targets": {
                    "text": {
                        "font_family": "Georgia",
                        "bold": True,
                        "italic": False,
                        "strikethrough": True,
                        "fill_color": "#ff0000",
                    }
                }
            },
            rich_text={
                "targets": {
                    "text": {
                        "doc": {
                            "type": "doc",
                            "content": [
                                {
                                    "type": "paragraph",
                                    "content": [
                                        {
                                            "type": "text",
                                            "text": "Белая метка",
                                            "marks": [
                                                {"type": "italic"},
                                                {
                                                    "type": "textStyle",
                                                    "attrs": {"fontFamily": "Arial"},
                                                },
                                                {
                                                    "type": "highlight",
                                                    "attrs": {"color": "#ffffff"},
                                                },
                                            ],
                                        },
                                        {"type": "text", "text": " и база"},
                                    ],
                                }
                            ],
                        }
                    }
                }
            },
        )
    )

    paragraph = _nonempty_paragraphs(
        Document(render_scenario_docx(snapshot)).tables[0].rows[4].cells[2]
    )[0]
    marked_run, base_run = paragraph.runs
    assert marked_run.font.name == "Arial"
    assert marked_run.bold is True
    assert marked_run.italic is True
    assert marked_run.font.strike is True
    assert _shading_fill(marked_run._r.get_or_add_rPr()) is None
    assert base_run.font.name == "Georgia"
    assert base_run.bold is True
    assert base_run.italic is False
    assert base_run.font.strike is True
    assert _shading_fill(base_run._r.get_or_add_rPr()) == "FF0000"


def test_stale_or_invalid_rich_doc_falls_back_to_canonical_target_text() -> None:
    snapshot = _snapshot(
        _row(
            "snh",
            "Каноническая речь",
            speaker_text="Каноническое имя\nКаноническая должность",
            rich_text={
                "targets": {
                    "speaker_fio": {
                        "doc": {
                            "type": "doc",
                            "content": [
                                {
                                    "type": "paragraph",
                                    "content": [
                                        {
                                            "type": "text",
                                            "text": "Устаревшее имя",
                                            "marks": [{"type": "strike"}],
                                        }
                                    ],
                                }
                            ],
                        }
                    },
                    "speaker_position": {
                        "doc": {"type": "not-a-doc", "content": []}
                    },
                    "text": {
                        "doc": {
                            "type": "doc",
                            "content": [
                                {
                                    "type": "paragraph",
                                    "content": [
                                        {"type": "text", "text": "Каноническая речь"}
                                    ],
                                }
                            ],
                        }
                    },
                }
            },
        )
    )

    document = Document(render_scenario_docx(snapshot))
    paragraphs = _nonempty_paragraphs(document.tables[0].rows[4].cells[0])
    assert [paragraph.text for paragraph in paragraphs] == [
        "Каноническое имя",
        "Каноническая должность",
        "Каноническая речь",
    ]
    assert paragraphs[0].runs[0].font.strike is not True
    assert paragraphs[0].runs[0].bold is True
    assert paragraphs[0].runs[0].italic is True


def test_long_multiline_text_has_no_fixed_row_height() -> None:
    long_text = "\n".join(f"Синтетическая строка {index:03d}" for index in range(240))
    document = Document(render_scenario_docx(_snapshot(_row("zk", long_text))))
    body_row = document.tables[0].rows[4]

    assert len(_nonempty_paragraphs(body_row.cells[0])) == 240
    assert body_row._tr.find("./w:trPr/w:trHeight", body_row._tr.nsmap) is None


@pytest.mark.parametrize(
    "target",
    [
        "title",
        "rubric",
        "duration",
        "body-rich-run",
        "speaker",
        "additional-comment",
        "bundle-filename",
        "bundle-timecode",
    ],
)
def test_renderer_replaces_xml_invalid_c0_controls_at_every_text_boundary(
    target: str,
) -> None:
    invalid_c0 = "".join(
        chr(code)
        for code in range(0x20)
        if code not in {0x09, 0x0A, 0x0D}
    )
    unsafe = f"До{invalid_c0}После"
    safe = f"До{'�' * 29}После"
    body_text = unsafe if target == "body-rich-run" else "Основной текст"
    speaker_fio = unsafe if target == "speaker" else "Синтетический спикер"
    additional_comment = unsafe if target == "additional-comment" else ""
    bundle = DocxFileBundle(
        unsafe if target == "bundle-filename" else "",
        unsafe if target == "bundle-timecode" else "",
        "",
    )
    rich_text = (
        {
            "targets": {
                "text": {
                    "doc": {
                        "type": "doc",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": unsafe,
                                        "marks": [{"type": "bold"}],
                                    }
                                ],
                            }
                        ],
                    }
                }
            }
        }
        if target == "body-rich-run"
        else {}
    )
    snapshot = ScenarioDocxSnapshot(
        story_id=73,
        title=unsafe if target == "title" else "Синтетический выпуск",
        rubric_id=9,
        rubric_name=unsafe if target == "rubric" else "Учебная рубрика",
        duration_text=unsafe if target == "duration" else "12:34",
        revision=4,
        rows=(
            _row(
                "snh",
                body_text,
                speaker_text=f"{speaker_fio}\nСинтетическая должность",
                additional_comment=additional_comment,
                rich_text=rich_text,
                file_bundles=(bundle,) if bundle.file_name or bundle.tc_in else (),
            ),
        ),
    )
    original_snapshot = deepcopy(snapshot)

    document = Document(render_scenario_docx(snapshot))
    table = document.tables[0]
    body = table.rows[4]
    if target == "title":
        rendered = table.rows[0].cells[0].text
    elif target == "rubric":
        rendered = table.rows[1].cells[0].text
    elif target == "duration":
        rendered = table.rows[2].cells[0].text
    elif target == "body-rich-run":
        rendered = _nonempty_paragraphs(body.cells[0])[-1].text
    elif target == "speaker":
        rendered = _nonempty_paragraphs(body.cells[0])[0].text
    else:
        rendered = _nonempty_paragraphs(body.cells[1])[0].text

    assert rendered == (f"Хронометраж {safe}" if target == "duration" else safe)
    assert snapshot == original_snapshot


def test_renderer_keeps_package_private_and_does_not_touch_filesystem(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    watched_temp = tmp_path / "watched-temp"
    watched_storage = tmp_path / "application-storage"
    watched_temp.mkdir()
    watched_storage.mkdir()
    marker = watched_storage / "keep.txt"
    marker.write_text("unchanged", encoding="utf-8")
    monkeypatch.setenv("TMPDIR", str(watched_temp))
    monkeypatch.setattr(tempfile, "tempdir", str(watched_temp))
    before = (_tree_listing(watched_temp), _tree_listing(watched_storage))

    buffer = render_scenario_docx(_fixture_snapshot())

    after = (_tree_listing(watched_temp), _tree_listing(watched_storage))
    assert after == before
    assert isinstance(buffer, BytesIO)
    with ZipFile(buffer) as archive:
        names = set(archive.namelist())
        core_xml = archive.read("docProps/core.xml").decode("utf-8")
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert "word/comments.xml" not in names
    assert "python-docx" not in core_xml.casefold()
    assert "pavel" not in core_xml.casefold()
    assert "lastModifiedBy" not in core_xml
    assert "/private/" not in core_xml
    assert "first-reference.mov" not in core_xml
    assert "Отдельный синтетический комментарий" not in core_xml
    assert all(
        token not in document_xml
        for token in (
            "<w:commentRangeStart",
            "<w:commentReference",
            "<w:ins>",
            "<w:ins ",
            "<w:del>",
            "<w:del ",
        )
    )


@pytest.mark.parametrize(
    ("title", "story_id", "expected_utf8"),
    [
        ("  Новости / день .. * ?  ", 7, "Новости-день.docx"),
        ("CON", 8, "Сценарий-8.docx"),
        ("lPt9. ", 9, "Сценарий-9.docx"),
        ("\x00\n\t", 10, "Сценарий-10.docx"),
        ("А" * 121, 11, f"{'А' * 120}.docx"),
        ("Безопасное имя", 12, "Безопасное имя.docx"),
    ],
)
def test_safe_docx_filename_sanitizes_utf8_name(
    title: str,
    story_id: int,
    expected_utf8: str,
) -> None:
    assert safe_docx_filename(title, story_id) == (
        f"Scenario-{story_id}.docx",
        expected_utf8,
    )
