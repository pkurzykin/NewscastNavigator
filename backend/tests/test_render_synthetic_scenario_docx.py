from __future__ import annotations

import ast
from dataclasses import FrozenInstanceError
from io import BytesIO
from pathlib import Path
import subprocess
import sys
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
import pytest


REPO_ROOT = Path(__file__).resolve().parents[2]
SCRIPT = REPO_ROOT / "backend/scripts/render_synthetic_scenario_docx.py"


def _run_script(*args: str) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(SCRIPT), *args],
        cwd=REPO_ROOT / "backend",
        text=True,
        capture_output=True,
        check=False,
    )


def test_synthetic_script_requires_explicit_safe_docx_output(tmp_path: Path) -> None:
    help_result = _run_script("--help")
    missing = _run_script()
    wrong_suffix = _run_script("--output", str(tmp_path / "scenario.zip"))
    forbidden_parent = REPO_ROOT / "backend/.task9-forbidden-output/scenario.docx"
    forbidden = _run_script("--output", str(forbidden_parent))
    symlink_target = tmp_path / "target.docx"
    symlink_target.write_bytes(b"keep")
    symlink_output = tmp_path / "linked.docx"
    symlink_output.symlink_to(symlink_target)
    symlink = _run_script("--output", str(symlink_output))

    assert help_result.returncode == 0
    assert "Создать синтетический DOCX-макет" in help_result.stdout
    assert missing.returncode == 2
    assert "--output" in missing.stderr
    assert wrong_suffix.returncode != 0
    assert ".docx" in wrong_suffix.stderr
    assert not (tmp_path / "scenario.zip").exists()
    assert forbidden.returncode != 0
    assert "artifacts/product-reset/V1_1_0/docx-export" in forbidden.stderr
    assert not forbidden_parent.parent.exists()
    assert symlink.returncode != 0
    assert "символическую ссылку" in symlink.stderr
    assert symlink_target.read_bytes() == b"keep"


def test_synthetic_script_atomically_replaces_a_raced_output_symlink(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from scripts import render_synthetic_scenario_docx as script

    output = tmp_path / "synthetic-scenario.docx"
    redirected = tmp_path / "redirected.docx"
    redirected.write_bytes(b"keep")

    def render_after_validation(_snapshot: object) -> BytesIO:
        output.symlink_to(redirected)
        return BytesIO(b"synthetic-docx")

    monkeypatch.setattr(script, "render_scenario_docx", render_after_validation)
    monkeypatch.setattr(
        sys,
        "argv",
        [str(SCRIPT), "--output", str(output)],
    )

    assert script.main() == 0
    assert output.is_file()
    assert not output.is_symlink()
    assert output.read_bytes() == b"synthetic-docx"
    assert redirected.read_bytes() == b"keep"


def test_atomic_output_removes_its_temporary_file_when_replace_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from scripts import render_synthetic_scenario_docx as script

    output = tmp_path / "synthetic-scenario.docx"

    def fail_replace(_source: object, _destination: object) -> None:
        raise OSError("synthetic replace failure")

    monkeypatch.setattr(script.os, "replace", fail_replace)

    with pytest.raises(OSError, match="synthetic replace failure"):
        script._write_output_atomically(output, b"synthetic-docx")

    assert not output.exists()
    assert tuple(tmp_path.iterdir()) == ()


def test_synthetic_script_renders_reopenable_five_block_multipage_fixture(
    tmp_path: Path,
) -> None:
    output = tmp_path / "nested/synthetic-scenario.docx"

    result = _run_script("--output", str(output))

    assert result.returncode == 0, result.stderr
    assert result.stdout == ""
    assert output.is_file()
    document = Document(output)
    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.rows) == 9
    all_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    for marker in (
        "СИНТЕТИЧЕСКИЙ МАКЕТ 1.1.0",
        "СИНТЕТИЧЕСКАЯ РУБРИКА",
        "Хронометраж до 02:15",
        "БЛОК-ПОДВОДКА",
        "БЛОК-ЗК",
        "БЛОК-ЗК-ГЕО",
        "БЛОК-СНХ",
        "БЛОК-ЛАЙФ",
        "СИНТЕТИЧЕСКАЯ СТРОКА 240",
        "synthetic-bundle-a.mov",
        "synthetic-bundle-b.mov",
    ):
        assert marker in all_text
    for unsafe_path_marker in (
        "/Volumes/",
        "C:\\Users\\",
    ):
        assert unsafe_path_marker not in all_text

    font_names: set[str] = set()
    fills: set[str] = set()
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    fonts = run._r.get_or_add_rPr().find(qn("w:rFonts"))
                    if fonts is not None:
                        name = fonts.get(qn("w:ascii"))
                        if name:
                            font_names.add(name)
                    shading = run._r.get_or_add_rPr().find(qn("w:shd"))
                    if shading is not None:
                        fill = shading.get(qn("w:fill"))
                        if fill:
                            fills.add(fill)
    assert font_names == {"PT Sans", "Arial", "Georgia", "Times New Roman", "Roboto Slab"}
    assert {"FFFF00", "FF0000", "00FF00", "0000FF", "FFA500"} <= fills
    with ZipFile(output) as archive:
        core_xml = archive.read("docProps/core.xml").decode("utf-8")
    assert "lastModifiedBy" not in core_xml
    assert str(tmp_path) not in core_xml


def test_synthetic_snapshot_has_frozen_empty_and_nonempty_duration_variants() -> None:
    from scripts.render_synthetic_scenario_docx import build_synthetic_snapshot

    populated = build_synthetic_snapshot(duration_text="до 02:15")
    empty = build_synthetic_snapshot(duration_text=None)

    assert populated.duration_text == "до 02:15"
    assert empty.duration_text is None
    assert len(populated.rows) == 5
    with pytest.raises(FrozenInstanceError):
        populated.title = "изменение запрещено"  # type: ignore[misc]
    with pytest.raises(TypeError):
        populated.rows[0].formatting["targets"] = {}  # type: ignore[index]


def test_runtime_route_never_imports_synthetic_render_script() -> None:
    route = REPO_ROOT / "backend/app/api/routes/scenario.py"
    tree = ast.parse(route.read_text(encoding="utf-8"))
    imported_modules = {
        alias.name
        for node in ast.walk(tree)
        if isinstance(node, ast.Import)
        for alias in node.names
    } | {
        node.module or ""
        for node in ast.walk(tree)
        if isinstance(node, ast.ImportFrom)
    }

    assert "scripts.render_synthetic_scenario_docx" not in imported_modules
