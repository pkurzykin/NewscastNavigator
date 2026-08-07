from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
import tempfile
from urllib.parse import unquote

from docx import Document
import pytest
from sqlalchemy import func, select, text

from app.core.security import hash_password
from app.db.base import Base
from app.db.models import (
    Notification,
    Rubric,
    Scenario,
    ScenarioRow,
    Story,
    StoryEvent,
    StoryProductionState,
    StoryWorkflowState,
    User,
    UserFunction,
)
from app.db.session import SessionLocal, engine


PASSWORD = "Scenario-Export-2026!"
DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _create_user(username: str, *, functions: tuple[str, ...] = ()) -> int:
    with SessionLocal() as db:
        user = User(
            username=username,
            display_name=f"Синтетический пользователь {username}",
            position="Сотрудник",
            password_hash=hash_password(PASSWORD),
            is_active=True,
            must_change_password=False,
            functions=[UserFunction(function_code=code) for code in functions],
        )
        db.add(user)
        db.commit()
        return user.id


def _create_story(
    *,
    author_user_id: int,
    title: str = "Синтетический экспорт",
    duration_text: str | None = "03:45",
    archived: bool = False,
) -> tuple[int, dict[str, object]]:
    now = datetime(2026, 8, 6, 9, 30, tzinfo=UTC)
    with SessionLocal() as db:
        rubric = Rubric(name=f"Рубрика {title}", is_active=True)
        db.add(rubric)
        db.flush()
        story = Story(
            title=title,
            rubric_id=rubric.id,
            author_user_id=author_user_id,
            duration_text=duration_text,
            aired_at=now if archived else None,
            aired_by_user_id=author_user_id if archived else None,
            archived_at=now if archived else None,
            archived_by_user_id=author_user_id if archived else None,
        )
        db.add(story)
        db.flush()
        scenario = Scenario(story_id=story.id, revision_no=4)
        workflow = StoryWorkflowState(
            story_id=story.id,
            editorial_revision=4,
            editorial_by_user_id=author_user_id,
            editorial_at=now,
        )
        production = StoryProductionState(
            story_id=story.id,
            voiceover_ready=True,
            voiceover_ready_by_user_id=author_user_id,
            voiceover_ready_at=now,
        )
        db.add_all([scenario, workflow, production])
        db.flush()
        db.add(
            ScenarioRow(
                scenario_id=scenario.id,
                segment_uid=f"seg-export-{story.id}",
                order_index=1,
                block_type="zk",
                text="Текст синтетического экспорта",
            )
        )
        db.add(
            StoryEvent(
                story_id=story.id,
                event_code="synthetic_export_baseline",
                actor_user_id=author_user_id,
                revision_no=4,
                payload={"source": "test"},
            )
        )
        db.add(
            Notification(
                recipient_user_id=author_user_id,
                story_id=story.id,
                kind="synthetic_export_baseline",
                actor_user_id=author_user_id,
                payload={"source": "test"},
            )
        )
        db.commit()
        return story.id, {
            "expected_revision": 4,
            "expected_title": title,
            "expected_rubric_id": rubric.id,
            "expected_duration_text": duration_text,
        }


def _login(client, username: str) -> dict[str, str]:
    response = client.post(
        "/api/v1/auth/login",
        json={"username": username, "password": PASSWORD},
    )
    assert response.status_code == 200, response.text
    return dict(response.cookies)


def _model_rows(
    db,
    model: object,
    *,
    story_id: int,
) -> tuple[tuple[tuple[str, object], ...], ...]:
    rows = db.execute(
        select(model)
        .where(getattr(model, "story_id") == story_id)
        .order_by(getattr(model, "id"))
    ).scalars()
    return tuple(
        tuple((column.name, getattr(row, column.name)) for column in model.__table__.columns)
        for row in rows
    )


def _aggregate_evidence(story_id: int) -> dict[str, object]:
    with SessionLocal() as db:
        story = db.get(Story, story_id)
        scenario = db.scalar(select(Scenario).where(Scenario.story_id == story_id))
        workflow = db.get(StoryWorkflowState, story_id)
        production = db.get(StoryProductionState, story_id)
        assert story is not None and scenario is not None
        assert workflow is not None and production is not None
        return {
            "story_updated_at": story.updated_at,
            "scenario_revision": scenario.revision_no,
            "events": _model_rows(db, StoryEvent, story_id=story_id),
            "workflow": tuple(
                (column.name, getattr(workflow, column.name))
                for column in StoryWorkflowState.__table__.columns
            ),
            "production": tuple(
                (column.name, getattr(production, column.name))
                for column in StoryProductionState.__table__.columns
            ),
            "notifications": _model_rows(db, Notification, story_id=story_id),
            "all_table_counts": tuple(
                (table.name, db.scalar(select(func.count()).select_from(table)))
                for table in Base.metadata.sorted_tables
            ),
        }


def _tree_listing(path: Path) -> tuple[str, ...]:
    return tuple(sorted(str(item.relative_to(path)) for item in path.rglob("*")))


def test_export_requires_browser_session_and_rejects_captionpanels_bearer(client) -> None:
    author_id = _create_user("export-auth")
    story_id, payload = _create_story(author_user_id=author_id)

    unauthenticated = client.post(
        f"/api/v1/stories/{story_id}/scenario/export-docx",
        json=payload,
    )
    captionpanels_login = client.post(
        "/api/v1/auth/login",
        headers={"Origin": "null"},
        json={"username": "export-auth", "password": PASSWORD},
    )
    assert captionpanels_login.status_code == 200, captionpanels_login.text
    client.cookies.clear()
    bearer_only = client.post(
        f"/api/v1/stories/{story_id}/scenario/export-docx",
        headers={
            "Origin": "null",
            "Authorization": f"Bearer {captionpanels_login.json()['access_token']}",
        },
        json=payload,
    )

    for response in (unauthenticated, bearer_only):
        assert response.status_code == 401
        assert response.json()["error"]["code"] == "AUTH_REQUIRED"


def test_export_returns_story_not_found_for_authenticated_user(client) -> None:
    _create_user("export-missing")

    response = client.post(
        "/api/v1/stories/404404/scenario/export-docx",
        cookies=_login(client, "export-missing"),
        json={
            "expected_revision": 0,
            "expected_title": "Несуществующий синтетический сюжет",
            "expected_rubric_id": None,
            "expected_duration_text": None,
        },
    )

    assert response.status_code == 404
    assert response.json() == {
        "error": {
            "code": "STORY_NOT_FOUND",
            "message": "Сюжет не найден",
            "details": {},
        }
    }


def test_any_authenticated_user_exports_active_and_archived_scenarios(client) -> None:
    author_id = _create_user("export-author")
    _create_user("export-reader")
    active_id, active_payload = _create_story(author_user_id=author_id)
    archived_id, archived_payload = _create_story(
        author_user_id=author_id,
        title="Архивный синтетический экспорт",
        archived=True,
    )

    for username in ("export-author", "export-reader"):
        cookies = _login(client, username)
        for story_id, payload in (
            (active_id, active_payload),
            (archived_id, archived_payload),
        ):
            response = client.post(
                f"/api/v1/stories/{story_id}/scenario/export-docx",
                cookies=cookies,
                json=payload,
            )
            assert response.status_code == 200, response.text
            assert response.content.startswith(b"PK")


@pytest.mark.parametrize(
    "payload",
    [
        {},
        {"expected_title": "Без редакции"},
        {"expected_revision": 4},
        {"expected_revision": -1, "expected_title": "Отрицательная редакция"},
        {"expected_revision": 4, "expected_title": "Т" * 256},
        {
            "expected_revision": 4,
            "expected_title": "Слишком длинный хронометраж",
            "expected_duration_text": "1" * 65,
        },
        {
            "expected_revision": 4,
            "expected_title": "Неверная рубрика",
            "expected_rubric_id": 0,
        },
    ],
    ids=[
        "all-required-missing",
        "revision-missing",
        "title-missing",
        "negative-revision",
        "title-too-long",
        "duration-too-long",
        "rubric-id-too-small",
    ],
)
def test_export_validates_required_fields_and_exact_bounds(client, payload) -> None:
    author_id = _create_user("export-validation")
    story_id, _matching_payload = _create_story(author_user_id=author_id)

    response = client.post(
        f"/api/v1/stories/{story_id}/scenario/export-docx",
        cookies=_login(client, "export-validation"),
        json=payload,
    )

    assert response.status_code == 422
    assert response.json()["error"]["code"] == "VALIDATION_ERROR"


def test_export_accepts_exact_title_and_duration_upper_bounds(client) -> None:
    author_id = _create_user("export-upper-bounds")
    title = "Т" * 255
    duration_text = "1" * 64
    story_id, payload = _create_story(
        author_user_id=author_id,
        title=title,
        duration_text=duration_text,
    )

    assert payload["expected_title"] == title
    assert payload["expected_duration_text"] == duration_text
    response = client.post(
        f"/api/v1/stories/{story_id}/scenario/export-docx",
        cookies=_login(client, "export-upper-bounds"),
        json=payload,
    )

    assert response.status_code == 200, response.text
    assert response.headers["content-type"] == DOCX_CONTENT_TYPE
    assert response.content.startswith(b"PK")


def test_created_title_with_all_line_boundaries_is_canonical_and_exportable(client) -> None:
    _create_user("export-title-create", functions=("author",))
    cookies = _login(client, "export-title-create")
    with SessionLocal() as db:
        rubric = Rubric(name="Синтетическая рубрика создания", is_active=True)
        db.add(rubric)
        db.commit()
        rubric_id = rubric.id
    raw_title = "  Создание\r\nчерез CRLF\rчерез CR\nчерез LF  "
    canonical_title = "Создание через CRLF через CR через LF"

    created = client.post(
        "/api/v1/stories",
        cookies=cookies,
        json={"title": raw_title, "rubric_id": rubric_id},
    )
    assert created.status_code == 200, created.text
    story_id = created.json()["resource"]["id"]
    scenario = client.get(
        f"/api/v1/stories/{story_id}/scenario",
        cookies=cookies,
    )
    assert scenario.status_code == 200, scenario.text
    exported = client.post(
        f"/api/v1/stories/{story_id}/scenario/export-docx",
        cookies=cookies,
        json={
            "expected_revision": scenario.json()["scenario"]["revision"],
            "expected_title": raw_title,
            "expected_rubric_id": rubric_id,
            "expected_duration_text": None,
        },
    )

    assert exported.status_code == 200, exported.text
    assert scenario.json()["story"]["title"] == canonical_title
    assert Document(BytesIO(exported.content)).tables[0].rows[0].cells[0].paragraphs[0].text == (
        canonical_title
    )


def test_metadata_title_with_all_line_boundaries_is_canonical_and_exportable(client) -> None:
    author_id = _create_user("export-title-update")
    story_id, payload = _create_story(author_user_id=author_id)
    cookies = _login(client, "export-title-update")
    raw_title = "  Правка\r\nчерез CRLF\rчерез CR\nчерез LF  "
    canonical_title = "Правка через CRLF через CR через LF"

    updated = client.patch(
        f"/api/v1/stories/{story_id}/metadata",
        cookies=cookies,
        json={"title": raw_title},
    )
    assert updated.status_code == 200, updated.text
    scenario = client.get(
        f"/api/v1/stories/{story_id}/scenario",
        cookies=cookies,
    )
    assert scenario.status_code == 200, scenario.text
    payload["expected_title"] = raw_title
    exported = client.post(
        f"/api/v1/stories/{story_id}/scenario/export-docx",
        cookies=cookies,
        json=payload,
    )

    assert exported.status_code == 200, exported.text
    assert scenario.json()["story"]["title"] == canonical_title
    assert Document(BytesIO(exported.content)).tables[0].rows[0].cells[0].paragraphs[0].text == (
        canonical_title
    )


def test_export_rejects_snapshot_mismatch_with_exact_conflict(client) -> None:
    author_id = _create_user("export-conflict")
    story_id, payload = _create_story(author_user_id=author_id)
    payload["expected_revision"] = 3

    response = client.post(
        f"/api/v1/stories/{story_id}/scenario/export-docx",
        cookies=_login(client, "export-conflict"),
        json=payload,
    )

    assert response.status_code == 409
    assert response.json() == {
        "error": {
            "code": "EXPORT_SNAPSHOT_MISMATCH",
            "message": "Сюжет изменился. Обновите карточку и повторите экспорт.",
            "details": {},
        }
    }


def test_export_rejects_corrupted_block_type_with_stable_domain_error(client) -> None:
    author_id = _create_user("export-corrupted-block")
    story_id, payload = _create_story(author_user_id=author_id)
    with engine.begin() as connection:
        connection.exec_driver_sql("PRAGMA ignore_check_constraints = ON")
        connection.execute(
            text(
                "UPDATE scenario_rows SET block_type = 'corrupted' "
                "WHERE scenario_id = (SELECT id FROM scenarios WHERE story_id = :story_id)"
            ),
            {"story_id": story_id},
        )
        connection.exec_driver_sql("PRAGMA ignore_check_constraints = OFF")

    response = client.post(
        f"/api/v1/stories/{story_id}/scenario/export-docx",
        cookies=_login(client, "export-corrupted-block"),
        json=payload,
    )

    assert response.status_code == 409
    assert response.json() == {
        "error": {
            "code": "EXPORT_UNSUPPORTED_BLOCK",
            "message": "Тип блока сценария не поддерживается для экспорта.",
            "details": {},
        }
    }


def test_export_returns_safe_exact_headers_and_reopenable_docx(client) -> None:
    author_id = _create_user("export-headers")
    story_id, payload = _create_story(
        author_user_id=author_id,
        title="Новости / день .. * ?",
    )
    payload["expected_title"] = "Новости\r\n/ день .. * ?"

    response = client.post(
        f"/api/v1/stories/{story_id}/scenario/export-docx",
        cookies=_login(client, "export-headers"),
        json=payload,
    )

    assert response.status_code == 200, response.text
    assert response.headers["content-type"] == DOCX_CONTENT_TYPE
    assert response.headers["cache-control"] == "no-store"
    disposition = response.headers["content-disposition"]
    assert disposition == (
        f'attachment; filename="Scenario-{story_id}.docx"; filename*=UTF-8\'\''
        "%D0%9D%D0%BE%D0%B2%D0%BE%D1%81%D1%82%D0%B8-"
        "%D0%B4%D0%B5%D0%BD%D1%8C.docx"
    )
    encoded_name = disposition.split("filename*=UTF-8''", 1)[1]
    assert unquote(encoded_name) == "Новости-день.docx"
    assert "\r" not in disposition and "\n" not in disposition
    assert all(
        character not in unquote(encoded_name)
        for character in ('/', '\\', ':', '*', '?', '"', '<', '>', '|')
    )
    assert len(response.content) > 1_000
    assert response.content.startswith(b"PK")
    document = Document(BytesIO(response.content))
    assert document.tables[0].rows[0].cells[0].paragraphs[0].text == "Новости / день .. * ?"
    assert (
        "Текст синтетического экспорта"
        in document.tables[0].rows[3].cells[0].text
    )


def test_duplicate_exports_do_not_mutate_database_or_create_files(
    client,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    author_id = _create_user("export-side-effects")
    story_id, payload = _create_story(author_user_id=author_id)
    cookies = _login(client, "export-side-effects")
    watched_temp = tmp_path / "watched-temp"
    watched_storage = tmp_path / "application-storage"
    watched_temp.mkdir()
    watched_storage.mkdir()
    (watched_storage / "keep.txt").write_text("unchanged", encoding="utf-8")
    monkeypatch.setenv("TMPDIR", str(watched_temp))
    monkeypatch.setattr(tempfile, "tempdir", str(watched_temp))
    before_state = _aggregate_evidence(story_id)
    before_files = (_tree_listing(watched_temp), _tree_listing(watched_storage))

    first = client.post(
        f"/api/v1/stories/{story_id}/scenario/export-docx",
        cookies=cookies,
        json=payload,
    )
    second = client.post(
        f"/api/v1/stories/{story_id}/scenario/export-docx",
        cookies=cookies,
        json=payload,
    )

    assert first.status_code == second.status_code == 200
    assert Document(BytesIO(first.content)).tables[0].rows[3].cells[0].text
    assert Document(BytesIO(second.content)).tables[0].rows[3].cells[0].text
    assert _aggregate_evidence(story_id) == before_state
    assert (_tree_listing(watched_temp), _tree_listing(watched_storage)) == before_files
