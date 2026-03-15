# app.py
from __future__ import annotations

import html
import io
import json
import os
import re
import sqlite3
import textwrap
import time
import hmac
import hashlib
import base64
from datetime import date, datetime
from pathlib import Path
from typing import Any

import streamlit as st

try:
    import pandas as pd
except Exception:
    pd = None

try:
    from st_aggrid import AgGrid, DataReturnMode, GridOptionsBuilder, GridUpdateMode
except Exception:
    AgGrid = None
    DataReturnMode = None
    GridOptionsBuilder = None
    GridUpdateMode = None

from auth import authenticate
from db import get_conn, init_db
from permissions import can_edit

st.set_page_config(page_title="Newscast Navigator", layout="wide")

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas as pdf_canvas
except Exception:
    A4 = None
    TTFont = None
    pdf_canvas = None
    pdfmetrics = None


STATUS_VALUES = [
    "draft",
    "reviewed",
    "in_editing",
    "in_proofreading",
    "ready",
    "delivered",
    "archived",
]

STATUS_LABELS = {
    "draft": "Черновик",
    "reviewed": "На проверке",
    "in_editing": "В работе",
    "in_proofreading": "На корректуре",
    "ready": "Готово",
    "delivered": "Сдано",
    "archived": "Архив",
}

BLOCK_TYPE_OPTIONS = [
    ("podvodka", "Подводка"),
    ("zk", "ЗК"),
    ("life", "Лайф"),
    ("snh", "СНХ"),
]
BLOCK_TYPE_LABELS = [label for _, label in BLOCK_TYPE_OPTIONS]

PROJECT_META_EDIT_ROLES = ("admin", "editor", "author")
PROJECT_ASSIGN_EDIT_ROLES = ("admin", "editor")
APP_ROOT = Path(__file__).resolve().parent
DEFAULT_STORAGE_ROOT = APP_ROOT / "storage"
ALLOWED_UPLOAD_EXTENSIONS = {
    ".mp4",
    ".mov",
    ".mxf",
    ".mp3",
    ".wav",
    ".m4a",
    ".aac",
    ".jpg",
    ".jpeg",
    ".png",
    ".webp",
    ".pdf",
    ".docx",
    ".txt",
}
MAX_UPLOAD_SIZE_MB = 512
MAX_UPLOAD_SIZE_BYTES = MAX_UPLOAD_SIZE_MB * 1024 * 1024
DOCX_EXPORT_AVAILABLE = Document is not None
PDF_EXPORT_AVAILABLE = A4 is not None and TTFont is not None and pdf_canvas is not None and pdfmetrics is not None
PDF_FONT_NAME = "Helvetica"
PDF_FONT_REGISTERED = False
SESSION_TOKEN_TTL_SECONDS = 7 * 24 * 60 * 60


def iso_now() -> str:
    return datetime.now().isoformat(timespec="seconds")


def status_label(status_value: str) -> str:
    return STATUS_LABELS.get(status_value, status_value)


def format_date_only(iso_value: str | None) -> str:
    if not iso_value:
        return "-"
    try:
        return datetime.fromisoformat(iso_value).strftime("%d.%m.%Y")
    except ValueError:
        return str(iso_value)[:10]


def format_datetime(iso_value: str | None) -> str:
    if not iso_value:
        return "-"
    try:
        return datetime.fromisoformat(iso_value).strftime("%d.%m.%Y %H:%M")
    except ValueError:
        return str(iso_value)


def block_code_to_name(code: str) -> str:
    for block_code, label in BLOCK_TYPE_OPTIONS:
        if block_code == code:
            return label
    return code


def block_name_to_code(name: str) -> str:
    for block_code, label in BLOCK_TYPE_OPTIONS:
        if label == name:
            return block_code
    return "zk"


def parse_timecode_to_seconds(value: str) -> int | None:
    text = (value or "").strip()
    if not text:
        return None

    # Supports MM:SS and HH:MM:SS.
    if not re.match(r"^\d{2}:\d{2}(:\d{2})?$", text):
        return None

    parts = [int(part) for part in text.split(":")]
    if len(parts) == 2:
        minutes, seconds = parts
        if seconds >= 60:
            return None
        return minutes * 60 + seconds

    hours, minutes, seconds = parts
    if minutes >= 60 or seconds >= 60:
        return None
    return hours * 3600 + minutes * 60 + seconds


def get_session_secret() -> str:
    secret = os.getenv("NEWSCAST_SESSION_SECRET", "").strip()
    if secret:
        return secret
    return "newscast-dev-insecure-secret-change-me"


def _sign_session_payload(payload_json: str) -> str:
    signature = hmac.new(
        get_session_secret().encode("utf-8"),
        payload_json.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()
    return signature


def create_session_token(user_id: int) -> str:
    payload = {
        "uid": int(user_id),
        "exp": int(time.time()) + SESSION_TOKEN_TTL_SECONDS,
    }
    payload_json = json.dumps(payload, separators=(",", ":"), sort_keys=True)
    signature = _sign_session_payload(payload_json)
    raw = f"{payload_json}.{signature}".encode("utf-8")
    return base64.urlsafe_b64encode(raw).decode("ascii")


def verify_session_token(token: str) -> int | None:
    if not token:
        return None
    try:
        decoded = base64.urlsafe_b64decode(token.encode("ascii")).decode("utf-8")
        payload_json, signature = decoded.rsplit(".", 1)
    except Exception:
        return None

    expected_signature = _sign_session_payload(payload_json)
    if not hmac.compare_digest(expected_signature, signature):
        return None

    try:
        payload = json.loads(payload_json)
        user_id = int(payload["uid"])
        exp = int(payload["exp"])
    except Exception:
        return None

    if exp < int(time.time()):
        return None

    return user_id


def read_auth_token_from_query() -> str:
    token = ""
    try:
        token = st.query_params.get("auth", "")
    except Exception:
        try:
            token = st.experimental_get_query_params().get("auth", "")
        except Exception:
            token = ""
    if isinstance(token, list):
        return str(token[0]) if token else ""
    return str(token or "")


def write_auth_token_to_query(token: str) -> None:
    try:
        st.query_params["auth"] = token
        return
    except Exception:
        pass
    try:
        st.experimental_set_query_params(auth=token)
    except Exception:
        pass


def clear_auth_token_from_query() -> None:
    try:
        del st.query_params["auth"]
        return
    except Exception:
        pass
    try:
        st.experimental_set_query_params()
    except Exception:
        pass


def fetch_user_by_id(cursor: sqlite3.Cursor, user_id: int) -> dict[str, Any] | None:
    row = cursor.execute(
        "SELECT id, username, role FROM users WHERE id = ?",
        (user_id,),
    ).fetchone()
    if not row:
        return None
    return {
        "id": int(row[0]),
        "username": row[1],
        "role": row[2],
    }


def fetch_users(cursor: sqlite3.Cursor) -> list[dict[str, Any]]:
    rows = cursor.execute(
        "SELECT id, username, role FROM users ORDER BY username"
    ).fetchall()
    return [{"id": int(row[0]), "username": row[1], "role": row[2]} for row in rows]


def user_label(users_map: dict[int, dict[str, Any]], user_id: int | None) -> str:
    if user_id is None:
        return "-"
    user_data = users_map.get(user_id)
    if not user_data:
        return "-"
    return f"{user_data['username']} ({user_data['role']})"


def option_index(options: list[Any], value: Any) -> int:
    if value in options:
        return options.index(value)
    return 0


def comments_have_project_scope(cursor: sqlite3.Cursor) -> bool:
    try:
        rows = cursor.execute("PRAGMA table_info(comments)").fetchall()
    except sqlite3.OperationalError:
        return False
    return any((row[1] or "") == "project_id" for row in rows)


def sanitize_file_name(file_name: str) -> str:
    base_name = os.path.basename(file_name or "").strip()
    if not base_name:
        return "upload.bin"

    sanitized = re.sub(r"[^A-Za-z0-9._-]+", "_", base_name)
    sanitized = sanitized.strip("._")
    if not sanitized:
        return "upload.bin"
    return sanitized


def get_storage_root_path() -> Path:
    root_from_env = os.getenv("NEWSCAST_STORAGE_ROOT", "").strip()
    if root_from_env:
        root = Path(root_from_env).expanduser()
        if not root.is_absolute():
            root = (APP_ROOT / root).resolve()
        return root
    return DEFAULT_STORAGE_ROOT


def resolve_project_storage_dir(project_id: int, project_file_root: str) -> Path:
    value = (project_file_root or "").strip()
    if value:
        base = Path(value).expanduser()
        if not base.is_absolute():
            base = (APP_ROOT / base).resolve()
        project_dir = base / f"project_{project_id}"
    else:
        project_dir = get_storage_root_path() / "projects" / str(project_id)

    project_dir.mkdir(parents=True, exist_ok=True)
    return project_dir


def validate_uploaded_file(uploaded_file: Any) -> tuple[bool, str]:
    if uploaded_file is None:
        return False, "Файл не выбран."

    extension = Path(uploaded_file.name or "").suffix.lower()
    if extension not in ALLOWED_UPLOAD_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_UPLOAD_EXTENSIONS))
        return False, f"Недопустимый тип файла: {extension or '(без расширения)'}. Разрешены: {allowed}"

    size_bytes = int(getattr(uploaded_file, "size", 0) or 0)
    if size_bytes <= 0:
        # Fallback for older Streamlit versions.
        try:
            size_bytes = len(uploaded_file.getbuffer())
        except Exception:
            size_bytes = 0
    if size_bytes <= 0:
        return False, "Не удалось определить размер файла."

    if size_bytes > MAX_UPLOAD_SIZE_BYTES:
        return (
            False,
            f"Слишком большой файл: {size_bytes // (1024 * 1024)} MB. Лимит: {MAX_UPLOAD_SIZE_MB} MB.",
        )

    return True, ""


def save_uploaded_file_to_disk(
    uploaded_file: Any,
    destination_dir: Path,
) -> tuple[Path, str, int]:
    destination_dir.mkdir(parents=True, exist_ok=True)
    safe_name = sanitize_file_name(uploaded_file.name or "upload.bin")
    time_prefix = datetime.now().strftime("%Y%m%d_%H%M%S")
    final_name = f"{time_prefix}_{safe_name}"
    destination_path = destination_dir / final_name

    counter = 1
    while destination_path.exists():
        final_name = f"{time_prefix}_{counter}_{safe_name}"
        destination_path = destination_dir / final_name
        counter += 1

    payload = uploaded_file.getvalue()
    with destination_path.open("wb") as target:
        target.write(payload)

    return destination_path, safe_name, len(payload)


def fetch_project_files(
    cursor: sqlite3.Cursor,
    project_id: int,
) -> list[dict[str, Any]]:
    rows = cursor.execute(
        """
        SELECT
            pf.id,
            pf.original_name,
            pf.storage_path,
            COALESCE(pf.mime_type, '') AS mime_type,
            COALESCE(pf.file_size, 0) AS file_size,
            pf.uploaded_at,
            COALESCE(u.username, '-') AS uploaded_by_name
        FROM project_files pf
        LEFT JOIN users u ON u.id = pf.uploaded_by
        WHERE pf.project_id = ? AND pf.element_id IS NULL
        ORDER BY pf.id DESC
        """,
        (project_id,),
    ).fetchall()
    return [
        {
            "id": int(row[0]),
            "original_name": row[1] or "",
            "storage_path": row[2] or "",
            "mime_type": row[3] or "",
            "file_size": int(row[4] or 0),
            "uploaded_at": row[5],
            "uploaded_by_name": row[6] or "-",
        }
        for row in rows
    ]


def add_project_file_record(
    conn: sqlite3.Connection,
    cursor: sqlite3.Cursor,
    *,
    project_id: int,
    element_id: int | None,
    original_name: str,
    storage_path: str,
    mime_type: str,
    file_size: int,
    uploaded_by: int,
) -> None:
    cursor.execute(
        """
        INSERT INTO project_files (
            project_id, element_id, original_name, storage_path,
            mime_type, file_size, uploaded_by, uploaded_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            project_id,
            element_id,
            original_name,
            storage_path,
            mime_type,
            file_size,
            uploaded_by,
            iso_now(),
        ),
    )
    if element_id is not None:
        cursor.execute(
            "UPDATE script_elements SET file_name = ? WHERE id = ?",
            (original_name, element_id),
        )
    log_project_event(
        cursor,
        project_id=project_id,
        event_type="file_uploaded",
        actor_user_id=uploaded_by,
        new_value=original_name,
    )
    conn.commit()


def remove_project_file(
    conn: sqlite3.Connection,
    cursor: sqlite3.Cursor,
    *,
    project_id: int,
    file_id: int,
    actor_user_id: int,
) -> tuple[bool, str]:
    row = cursor.execute(
        """
        SELECT id, element_id, original_name, storage_path
        FROM project_files
        WHERE id = ? AND project_id = ?
        """,
        (file_id, project_id),
    ).fetchone()
    if not row:
        return False, "Файл не найден."

    _, element_id, original_name, storage_path = row
    file_path = Path(storage_path)
    if file_path.exists():
        try:
            file_path.unlink()
        except OSError as exc:
            return False, f"Не удалось удалить файл на диске: {exc}"

    cursor.execute("DELETE FROM project_files WHERE id = ?", (file_id,))
    if element_id is not None:
        latest_file = cursor.execute(
            """
            SELECT original_name
            FROM project_files
            WHERE element_id = ?
            ORDER BY id DESC
            LIMIT 1
            """,
            (element_id,),
        ).fetchone()
        cursor.execute(
            "UPDATE script_elements SET file_name = ? WHERE id = ?",
            ((latest_file[0] if latest_file else ""), element_id),
        )
    log_project_event(
        cursor,
        project_id=project_id,
        event_type="file_deleted",
        actor_user_id=actor_user_id,
        old_value=original_name,
    )
    conn.commit()
    return True, "Файл удален."


def ensure_pdf_font() -> str:
    global PDF_FONT_REGISTERED, PDF_FONT_NAME

    if not PDF_EXPORT_AVAILABLE:
        return "Helvetica"
    if PDF_FONT_REGISTERED:
        return PDF_FONT_NAME

    candidates = [
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/usr/share/fonts/dejavu/DejaVuSans.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        Path("/Library/Fonts/Arial Unicode.ttf"),
    ]
    for candidate in candidates:
        if not candidate.exists():
            continue
        try:
            pdfmetrics.registerFont(TTFont("NewscastUnicode", str(candidate)))
            PDF_FONT_NAME = "NewscastUnicode"
            PDF_FONT_REGISTERED = True
            return PDF_FONT_NAME
        except Exception:
            continue

    PDF_FONT_REGISTERED = True
    PDF_FONT_NAME = "Helvetica"
    return PDF_FONT_NAME


def inject_global_ui_css() -> None:
    st.markdown(
        """
        <style>
        :root {
            color-scheme: light;
        }
        .stApp {
            background: #efefef;
            color: #111 !important;
        }
        [data-testid="stMainBlockContainer"] {
            max-width: 1680px;
            padding-top: 1rem;
        }
        [data-testid="stSidebar"] {
            background: #f4f4f4;
            color: #111 !important;
        }
        p, span, label, div, li {
            color: #111 !important;
        }
        input, textarea, [data-baseweb="select"] > div {
            color: #111 !important;
            background: #fff !important;
        }
        .stTextInput label, .stTextArea label, .stSelectbox label, .stMultiSelect label, .stDateInput label, .stCheckbox label {
            color: #111 !important;
            font-weight: 600;
        }
        .nn-screen-header {
            background: #f4ea52;
            border: 1px solid #d7cb38;
            text-align: center;
            font-weight: 800;
            font-size: 36px;
            line-height: 1.2;
            padding: 10px 12px;
            margin-bottom: 12px;
            letter-spacing: 0.2px;
        }
        .nn-screen-subtitle {
            font-size: 13px;
            color: #333;
            margin: 0 0 12px 0;
        }
        .nn-block-cards {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(190px, 1fr));
            gap: 6px;
            margin-bottom: 12px;
        }
        .nn-card {
            min-height: 110px;
            border: 2px solid #2f7d32;
            padding: 8px;
            font-size: 14px;
            font-weight: 700;
            line-height: 1.25;
            color: #111;
        }
        .nn-card-pink { background: #f49aa0; }
        .nn-card-orange { background: #ff8b00; }
        .nn-card-yellow { background: #ffd329; }
        .nn-card-green { background: #7fff57; }
        .nn-card-blue { background: #66c8ff; border-color: #2b4a6f; }
        .nn-card-purple { background: #8c2a73; color: #fff; border-color: #5a1a4a; }
        .nn-editor-preview {
            border: 1px solid #a7a7a7;
            background: #fff;
            margin: 10px 0 14px 0;
        }
        .nn-editor-preview-title {
            background: #c4c4c4;
            font-weight: 800;
            text-align: center;
            padding: 10px;
            border-bottom: 1px solid #9f9f9f;
            font-size: 20px;
        }
        .nn-editor-grid {
            display: grid;
            grid-template-columns: 1.2fr 2.8fr 1.4fr 1fr 0.8fr 0.9fr 1.2fr;
            gap: 0;
        }
        .nn-editor-cell {
            border-right: 1px solid #9c9c9c;
            border-bottom: 1px solid #9c9c9c;
            min-height: 66px;
            padding: 8px;
            font-size: 13px;
            line-height: 1.25;
        }
        .nn-editor-cell b {
            display: block;
            margin-bottom: 4px;
        }
        .nn-rubric { background: #f593d2; }
        .nn-news { background: #f593d2; }
        .nn-author { background: #f5eb63; }
        .nn-ann { background: #f5eb63; }
        .nn-mode { background: #78c5ed; }
        .nn-duration { background: #59e63f; }
        .nn-status { background: #a0a0a0; }
        .nn-status-chip {
            display: inline-block;
            padding: 2px 8px;
            border-radius: 999px;
            font-size: 12px;
            font-weight: 700;
            border: 1px solid rgba(0, 0, 0, 0.15);
            background: #efefef;
        }
        .nn-status-draft { background: #ffe7a6; }
        .nn-status-reviewed { background: #dce8ff; }
        .nn-status-in_editing { background: #c4f0ff; }
        .nn-status-in_proofreading { background: #ffd9d9; }
        .nn-status-ready { background: #d6ffd6; }
        .nn-status-delivered { background: #ddd; }
        .nn-status-archived { background: #ececec; }
        .nn-soft-note {
            font-size: 12px;
            color: #3d3d3d;
            margin-top: 4px;
        }
        [data-testid="stDataEditor"] {
            border: 2px solid #9f9f9f;
            background: #fff;
        }
        [data-testid="stDataEditor"] [role="columnheader"] {
            background: #e9e9e9 !important;
            color: #111 !important;
            font-weight: 700 !important;
        }
        [data-testid="stDataEditor"] [role="gridcell"] {
            border-right: 1px solid #b5b5b5 !important;
            border-bottom: 1px solid #b5b5b5 !important;
            background: #fff !important;
            color: #111 !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_screen_header(section_label: str, subtitle: str = "") -> None:
    st.markdown(
        f"<div class='nn-screen-header'>Newscast Navigator&nbsp;&nbsp;{html.escape(section_label)}</div>",
        unsafe_allow_html=True,
    )
    if subtitle:
        st.markdown(
            f"<div class='nn-screen-subtitle'>{html.escape(subtitle)}</div>",
            unsafe_allow_html=True,
        )


def render_main_action_blocks() -> None:
    st.markdown(
        """
        <div class="nn-block-cards">
          <div class="nn-card nn-card-pink">Создать новый пустой</div>
          <div class="nn-card nn-card-pink">Создать на основе последнего</div>
          <div class="nn-card nn-card-pink">Создать на основе выбранного</div>
          <div class="nn-card nn-card-yellow">Export (Word, PDF)</div>
          <div class="nn-card nn-card-orange">Отобразить архив</div>
          <div class="nn-card nn-card-green">Сдать в архив (с подтверждением)</div>
          <div class="nn-card nn-card-blue">Кнопка авторизации/профиля</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_archive_action_blocks() -> None:
    st.markdown(
        """
        <div class="nn-block-cards">
          <div class="nn-card nn-card-green">Вернуть в работу (с подтверждением)</div>
          <div class="nn-card nn-card-blue">Кнопка авторизации/профиля</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _value_is_empty(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, float) and value != value:
        return True
    return str(value).strip() == ""


def _coerce_int_or_none(value: Any) -> int | None:
    if _value_is_empty(value):
        return None
    try:
        return int(str(value).strip())
    except Exception:
        return None


def normalize_speaker_text_for_editor(value: str) -> str:
    return (value or "").replace("\r\n", "\n").strip()


def build_editor_dataframe(elements: list[sqlite3.Row | tuple[Any, ...]]) -> "pd.DataFrame | None":
    if pd is None:
        return None

    rows: list[dict[str, Any]] = []
    for row in elements:
        (
            el_id,
            idx,
            text_value,
            block_type,
            speaker_text,
            file_name,
            tc_in,
            tc_out,
            additional_comment,
        ) = row
        block_type_code = block_type or "zk"
        rows.append(
            {
                "ID": int(el_id),
                "№": int(idx),
                "Блок": block_code_to_name(block_type_code),
                "Текст": text_value or "",
                "Титр": normalize_speaker_text_for_editor(speaker_text or ""),
                "Имя файла": file_name or "",
                "TC IN": tc_in or "",
                "TC OUT": tc_out or "",
                "Другой коммент": additional_comment or "",
            }
        )

    if not rows:
        rows = [
            {
                "ID": "",
                "№": 1,
                "Блок": "ЗК",
                "Текст": "",
                "Титр": "",
                "Имя файла": "",
                "TC IN": "",
                "TC OUT": "",
                "Другой коммент": "",
            }
        ]

    return pd.DataFrame(rows)


def normalize_editor_rows(table_rows: list[dict[str, Any]]) -> tuple[list[dict[str, Any]], list[str]]:
    normalized_rows: list[dict[str, Any]] = []
    errors: list[str] = []
    next_index = 1

    for row in table_rows:
        row_id = _coerce_int_or_none(row.get("ID"))

        block_name = str(row.get("Блок") or "ЗК").strip()
        if block_name not in BLOCK_TYPE_LABELS:
            block_name = "ЗК"
        block_type_code = block_name_to_code(block_name)

        text_value = str(row.get("Текст") or "").strip()
        legacy_fio_value = str(row.get("Титр: ФИО") or "").strip()
        legacy_position_value = str(row.get("Титр: должность") or "").strip()
        combined_legacy_speaker = ""
        if legacy_fio_value or legacy_position_value:
            combined_legacy_speaker = (
                f"{legacy_fio_value}\n{legacy_position_value}".strip()
                if block_type_code == "snh"
                else (legacy_fio_value or legacy_position_value)
            )

        speaker_value = str(
            row.get("Титр")
            or row.get("Титры/ФИО")
            or combined_legacy_speaker
            or ""
        ).strip()
        speaker_lines = [
            line.strip()
            for line in speaker_value.replace("\r\n", "\n").split("\n")
            if line.strip()
        ]
        file_name_value = str(row.get("Имя файла") or "").strip()
        tc_in_value = str(row.get("TC IN") or "").strip()
        tc_out_value = str(row.get("TC OUT") or "").strip()
        add_comment_value = str(row.get("Другой коммент") or "").strip()

        if (
            row_id is None
            and not text_value
            and not speaker_value
            and not file_name_value
            and not tc_in_value
            and not tc_out_value
            and not add_comment_value
        ):
            continue

        tc_in_seconds = parse_timecode_to_seconds(tc_in_value) if tc_in_value else None
        tc_out_seconds = parse_timecode_to_seconds(tc_out_value) if tc_out_value else None

        if tc_in_value and tc_in_seconds is None:
            errors.append(
                f"Строка {next_index}: неверный формат TC IN. Используйте MM:SS или HH:MM:SS."
            )
        if tc_out_value and tc_out_seconds is None:
            errors.append(
                f"Строка {next_index}: неверный формат TC OUT. Используйте MM:SS или HH:MM:SS."
            )
        if (
            tc_in_seconds is not None
            and tc_out_seconds is not None
            and tc_out_seconds < tc_in_seconds
        ):
            errors.append(f"Строка {next_index}: TC OUT не может быть меньше TC IN.")
        if block_type_code == "snh" and speaker_value and len(speaker_lines) < 2:
            errors.append(
                f"Строка {next_index}: для блока СНХ в колонке 'Титр' укажите две строки: ФИО и должность."
            )

        normalized_rows.append(
            {
                "id": row_id,
                "order_index": next_index,
                "block_type": block_type_code,
                "text": text_value,
                "speaker_text": "\n".join(speaker_lines) if block_type_code == "snh" else speaker_value,
                "file_name": file_name_value,
                "tc_in": tc_in_value,
                "tc_out": tc_out_value,
                "additional_comment": add_comment_value,
            }
        )
        next_index += 1

    return normalized_rows, errors


def delete_script_element_dependencies(
    cursor: sqlite3.Cursor,
    *,
    element_ids: list[int],
) -> None:
    if not element_ids:
        return

    placeholders = ",".join("?" for _ in element_ids)
    file_rows = cursor.execute(
        f"SELECT storage_path FROM project_files WHERE element_id IN ({placeholders})",
        element_ids,
    ).fetchall()
    for file_row in file_rows:
        file_path = Path(file_row[0] or "")
        if file_path.exists():
            try:
                file_path.unlink()
            except OSError:
                pass

    cursor.execute(
        f"DELETE FROM project_files WHERE element_id IN ({placeholders})",
        element_ids,
    )
    cursor.execute(
        f"DELETE FROM comments WHERE element_id IN ({placeholders})",
        element_ids,
    )


def save_editor_rows(
    conn: sqlite3.Connection,
    cursor: sqlite3.Cursor,
    *,
    project_id: int,
    rows: list[dict[str, Any]],
) -> tuple[int, int, int]:
    existing_rows = cursor.execute(
        "SELECT id FROM script_elements WHERE project_id = ?",
        (project_id,),
    ).fetchall()
    existing_ids = {int(row[0]) for row in existing_rows}
    incoming_ids = {row["id"] for row in rows if row["id"] is not None}

    removed_ids = sorted(existing_ids - incoming_ids)
    if removed_ids:
        delete_script_element_dependencies(cursor, element_ids=removed_ids)
        placeholders = ",".join("?" for _ in removed_ids)
        cursor.execute(
            f"DELETE FROM script_elements WHERE id IN ({placeholders})",
            removed_ids,
        )

    updated_count = 0
    inserted_count = 0
    for row in rows:
        row_id = row["id"]
        if row_id is not None and row_id in existing_ids:
            cursor.execute(
                """
                UPDATE script_elements
                SET
                    order_index = ?,
                    text = ?,
                    element_type = ?,
                    block_type = ?,
                    speaker_text = ?,
                    file_name = ?,
                    tc_in = ?,
                    tc_out = ?,
                    additional_comment = ?
                WHERE id = ?
                """,
                (
                    row["order_index"],
                    row["text"],
                    row["block_type"],
                    row["block_type"],
                    row["speaker_text"],
                    row["file_name"],
                    row["tc_in"],
                    row["tc_out"],
                    row["additional_comment"],
                    row_id,
                ),
            )
            updated_count += 1
        else:
            cursor.execute(
                """
                INSERT INTO script_elements (
                    project_id, order_index, text, element_type, block_type,
                    speaker_text, file_name, tc_in, tc_out, additional_comment
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    project_id,
                    row["order_index"],
                    row["text"],
                    row["block_type"],
                    row["block_type"],
                    row["speaker_text"],
                    row["file_name"],
                    row["tc_in"],
                    row["tc_out"],
                    row["additional_comment"],
                ),
            )
            inserted_count += 1

    conn.commit()
    return updated_count, inserted_count, len(removed_ids)


def fetch_project_export_payload(cursor: sqlite3.Cursor, project_id: int) -> dict[str, Any]:
    project_row = cursor.execute(
        """
        SELECT
            p.id,
            p.title,
            p.status,
            COALESCE(NULLIF(p.rubric, ''), '-') AS rubric,
            COALESCE(NULLIF(p.planned_duration, ''), '-') AS planned_duration,
            COALESCE(au.username, '-') AS author_name,
            COALESCE(eu.username, '-') AS executor_name,
            COALESCE(pu.username, '-') AS proofreader_name,
            p.created_at
        FROM projects p
        LEFT JOIN users au ON au.id = COALESCE(p.author_user_id, p.author_id)
        LEFT JOIN users eu ON eu.id = p.executor_user_id
        LEFT JOIN users pu ON pu.id = p.proofreader_user_id
        WHERE p.id = ?
        """,
        (project_id,),
    ).fetchone()

    elements_rows = cursor.execute(
        """
        SELECT
            order_index,
            COALESCE(NULLIF(block_type, ''), 'zk') AS block_type,
            COALESCE(text, '') AS text,
            COALESCE(speaker_text, '') AS speaker_text,
            COALESCE(file_name, '') AS file_name,
            COALESCE(tc_in, '') AS tc_in,
            COALESCE(tc_out, '') AS tc_out,
            COALESCE(additional_comment, '') AS additional_comment
        FROM script_elements
        WHERE project_id = ?
        ORDER BY order_index
        """,
        (project_id,),
    ).fetchall()

    if not project_row:
        return {"project": None, "elements": []}

    project = {
        "id": int(project_row[0]),
        "title": project_row[1] or "",
        "status": project_row[2] or "",
        "rubric": project_row[3] or "-",
        "planned_duration": project_row[4] or "-",
        "author_name": project_row[5] or "-",
        "executor_name": project_row[6] or "-",
        "proofreader_name": project_row[7] or "-",
        "created_at": project_row[8],
    }

    elements = [
        {
            "order_index": int(row[0]),
            "block_type": row[1] or "zk",
            "text": row[2] or "",
            "speaker_text": row[3] or "",
            "file_name": row[4] or "",
            "tc_in": row[5] or "",
            "tc_out": row[6] or "",
            "additional_comment": row[7] or "",
        }
        for row in elements_rows
    ]
    return {"project": project, "elements": elements}


def build_export_lines(payload: dict[str, Any]) -> list[str]:
    project = payload.get("project")
    if not project:
        return ["Нет данных проекта для экспорта."]

    lines: list[str] = [
        f"Newscast Navigator | Проект #{project['id']}",
        f"Название: {project['title']}",
        f"Статус: {status_label(project['status'])}",
        f"Рубрика: {project['rubric']}",
        f"Хронометраж: {project['planned_duration']}",
        f"Автор: {project['author_name']}",
        f"Исполнитель: {project['executor_name']}",
        f"Корректор: {project['proofreader_name']}",
        f"Дата создания: {format_datetime(project['created_at'])}",
        "",
        "Сценарий:",
    ]

    for row in payload.get("elements", []):
        lines.append(
            f"{row['order_index']}. {block_code_to_name(row['block_type'])} | "
            f"TC {row['tc_in'] or '-'} - {row['tc_out'] or '-'}"
        )
        lines.append(f"Текст: {row['text']}")
        lines.append(f"Титры/ФИО: {row['speaker_text'] or '-'}")
        lines.append(f"Имя файла: {row['file_name'] or '-'}")
        lines.append(f"Комментарий: {row['additional_comment'] or '-'}")
        lines.append("")

    return lines


def generate_docx_bytes(payload: dict[str, Any]) -> tuple[bytes | None, str]:
    if not DOCX_EXPORT_AVAILABLE:
        return None, "Для DOCX установите пакет: pip install python-docx"

    project = payload.get("project")
    if not project:
        return None, "Нет данных проекта для экспорта."

    doc = Document()
    doc.add_heading(f"Newscast Navigator — {project['title']}", level=1)
    doc.add_paragraph(f"Проект ID: {project['id']}")
    doc.add_paragraph(f"Статус: {status_label(project['status'])}")
    doc.add_paragraph(f"Рубрика: {project['rubric']}")
    doc.add_paragraph(f"Хронометраж: {project['planned_duration']}")
    doc.add_paragraph(f"Автор: {project['author_name']}")
    doc.add_paragraph(f"Исполнитель: {project['executor_name']}")
    doc.add_paragraph(f"Корректор: {project['proofreader_name']}")
    doc.add_paragraph(f"Дата создания: {format_datetime(project['created_at'])}")

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["№", "Блок", "Текст", "Титры/ФИО", "Имя файла", "TC IN", "TC OUT", "Комментарий"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for row in payload.get("elements", []):
        cells = table.add_row().cells
        cells[0].text = str(row["order_index"])
        cells[1].text = block_code_to_name(row["block_type"])
        cells[2].text = row["text"]
        cells[3].text = row["speaker_text"]
        cells[4].text = row["file_name"]
        cells[5].text = row["tc_in"]
        cells[6].text = row["tc_out"]
        cells[7].text = row["additional_comment"]

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), ""


def generate_pdf_bytes(payload: dict[str, Any]) -> tuple[bytes | None, str]:
    if not PDF_EXPORT_AVAILABLE:
        return None, "Для PDF установите пакет: pip install reportlab"

    lines = build_export_lines(payload)
    font_name = ensure_pdf_font()

    buffer = io.BytesIO()
    pdf = pdf_canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin_x = 36
    margin_y = 40
    line_height = 13
    y = height - margin_y

    pdf.setFont(font_name, 10)
    for line in lines:
        wrapped_lines = textwrap.wrap(line, width=105) or [""]
        for wrapped in wrapped_lines:
            if y <= margin_y:
                pdf.showPage()
                pdf.setFont(font_name, 10)
                y = height - margin_y
            pdf.drawString(margin_x, y, wrapped)
            y -= line_height

    pdf.save()
    return buffer.getvalue(), ""


def build_print_html(payload: dict[str, Any]) -> str:
    project = payload.get("project")
    if not project:
        return "<p>Нет данных проекта для печати.</p>"

    rows_html: list[str] = []
    for row in payload.get("elements", []):
        rows_html.append(
            "<tr>"
            f"<td>{row['order_index']}</td>"
            f"<td>{html.escape(block_code_to_name(row['block_type']))}</td>"
            f"<td>{html.escape(row['text'])}</td>"
            f"<td>{html.escape(row['speaker_text'])}</td>"
            f"<td>{html.escape(row['file_name'])}</td>"
            f"<td>{html.escape(row['tc_in'])}</td>"
            f"<td>{html.escape(row['tc_out'])}</td>"
            f"<td>{html.escape(row['additional_comment'])}</td>"
            "</tr>"
        )

    return (
        """
        <style>
        .nn-print-meta {
            margin-bottom: 10px;
            font-size: 13px;
            line-height: 1.35;
        }
        .nn-print-table {
            width: 100%;
            border-collapse: collapse;
            font-size: 12px;
        }
        .nn-print-table th, .nn-print-table td {
            border: 1px solid #777;
            padding: 5px 6px;
            text-align: left;
            vertical-align: top;
        }
        .nn-print-table td:nth-child(4) {
            white-space: pre-line;
        }
        .nn-print-table th {
            background: #efefef;
        }
        </style>
        """
        + "<div class='nn-print-meta'>"
        + f"<b>Проект:</b> {html.escape(project['title'])}<br>"
        + f"<b>Статус:</b> {html.escape(status_label(project['status']))}<br>"
        + f"<b>Рубрика:</b> {html.escape(project['rubric'])}<br>"
        + f"<b>Хронометраж:</b> {html.escape(project['planned_duration'])}<br>"
        + f"<b>Автор:</b> {html.escape(project['author_name'])}<br>"
        + f"<b>Исполнитель:</b> {html.escape(project['executor_name'])}<br>"
        + f"<b>Корректор:</b> {html.escape(project['proofreader_name'])}<br>"
        + f"<b>Дата создания:</b> {html.escape(format_datetime(project['created_at']))}<br>"
        + "</div>"
        + "<table class='nn-print-table'>"
        + "<thead><tr>"
        + "<th>№</th><th>Блок</th><th>Текст</th><th>Титры/ФИО</th>"
        + "<th>Имя файла</th><th>TC IN</th><th>TC OUT</th><th>Комментарий</th>"
        + "</tr></thead>"
        + "<tbody>"
        + "".join(rows_html)
        + "</tbody></table>"
    )


def log_project_event(
    cursor: sqlite3.Cursor,
    *,
    project_id: int,
    event_type: str,
    actor_user_id: int,
    old_value: str | None = None,
    new_value: str | None = None,
    meta_json: str | None = None,
) -> None:
    try:
        cursor.execute(
            """
            INSERT INTO project_events (
                project_id, event_type, old_value, new_value,
                actor_user_id, created_at, meta_json
            ) VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                project_id,
                event_type,
                old_value,
                new_value,
                actor_user_id,
                iso_now(),
                meta_json,
            ),
        )
    except sqlite3.OperationalError:
        # If migrations were not applied for some reason, app should not crash.
        return


def insert_default_script_rows(cursor: sqlite3.Cursor, project_id: int) -> None:
    for index, (block_code, block_name) in enumerate(BLOCK_TYPE_OPTIONS, start=1):
        cursor.execute(
            """
            INSERT INTO script_elements (
                project_id, order_index, text, element_type, block_type
            ) VALUES (?, ?, ?, ?, ?)
            """,
            (project_id, index, f"{block_name}: ", block_code, block_code),
        )


def create_empty_project(
    conn: sqlite3.Connection,
    cursor: sqlite3.Cursor,
    user_id: int,
) -> int:
    now = iso_now()
    cursor.execute(
        """
        INSERT INTO projects (
            title, topic, status, author_id, author_user_id,
            created_at, status_changed_at, status_changed_by, is_archived
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, 0)
        """,
        (
            "Новый проект",
            "",
            "draft",
            user_id,
            user_id,
            now,
            now,
            user_id,
        ),
    )
    project_id = int(cursor.lastrowid)
    insert_default_script_rows(cursor, project_id)
    log_project_event(
        cursor,
        project_id=project_id,
        event_type="project_created",
        actor_user_id=user_id,
    )
    conn.commit()
    return project_id


def clone_project(
    conn: sqlite3.Connection,
    cursor: sqlite3.Cursor,
    source_project_id: int,
    actor_user_id: int,
) -> int | None:
    source = cursor.execute(
        """
        SELECT
            title,
            topic,
            rubric,
            planned_duration,
            executor_user_id,
            proofreader_user_id
        FROM projects
        WHERE id = ?
        """,
        (source_project_id,),
    ).fetchone()
    if not source:
        return None

    now = iso_now()
    source_title, topic, rubric, planned_duration, executor_user_id, proofreader_user_id = source
    new_title = f"{source_title} (копия)"

    cursor.execute(
        """
        INSERT INTO projects (
            title, topic, status, author_id, author_user_id,
            created_at, rubric, planned_duration, source_project_id,
            executor_user_id, proofreader_user_id,
            status_changed_at, status_changed_by, is_archived
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0)
        """,
        (
            new_title,
            topic,
            "draft",
            actor_user_id,
            actor_user_id,
            now,
            rubric or "",
            planned_duration or "",
            source_project_id,
            executor_user_id,
            proofreader_user_id,
            now,
            actor_user_id,
        ),
    )
    new_project_id = int(cursor.lastrowid)

    source_rows = cursor.execute(
        """
        SELECT
            order_index,
            text,
            element_type,
            block_type,
            speaker_text,
            file_name,
            tc_in,
            tc_out,
            additional_comment
        FROM script_elements
        WHERE project_id = ?
        ORDER BY order_index
        """,
        (source_project_id,),
    ).fetchall()

    if source_rows:
        for row in source_rows:
            cursor.execute(
                """
                INSERT INTO script_elements (
                    project_id, order_index, text, element_type, block_type,
                    speaker_text, file_name, tc_in, tc_out, additional_comment
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (new_project_id, *row),
            )
    else:
        insert_default_script_rows(cursor, new_project_id)

    log_project_event(
        cursor,
        project_id=new_project_id,
        event_type="project_cloned",
        actor_user_id=actor_user_id,
        old_value=str(source_project_id),
    )
    conn.commit()
    return new_project_id


def archive_project(
    conn: sqlite3.Connection,
    cursor: sqlite3.Cursor,
    project_id: int,
    actor_user_id: int,
) -> bool:
    row = cursor.execute(
        "SELECT status, COALESCE(is_archived, 0) FROM projects WHERE id = ?",
        (project_id,),
    ).fetchone()
    if not row:
        return False

    old_status, is_archived = row
    if int(is_archived) == 1:
        return False

    now = iso_now()
    cursor.execute(
        """
        UPDATE projects
        SET
            is_archived = 1,
            archived_at = ?,
            archived_by = ?,
            status = ?,
            status_changed_at = ?,
            status_changed_by = ?
        WHERE id = ?
        """,
        (now, actor_user_id, "archived", now, actor_user_id, project_id),
    )
    log_project_event(
        cursor,
        project_id=project_id,
        event_type="status_changed",
        actor_user_id=actor_user_id,
        old_value=old_status,
        new_value="archived",
    )
    log_project_event(
        cursor,
        project_id=project_id,
        event_type="project_archived",
        actor_user_id=actor_user_id,
    )
    conn.commit()
    return True


def fetch_projects_for_main(
    cursor: sqlite3.Cursor,
    *,
    include_archived: bool,
    status_filter: list[str],
    rubric_query: str,
    participant_query: str,
    use_date_filter: bool,
    date_from: date,
    date_to: date,
) -> list[dict[str, Any]]:
    clauses: list[str] = []
    params: list[Any] = []

    if not include_archived:
        clauses.append("COALESCE(p.is_archived, 0) = 0")

    if status_filter:
        placeholders = ",".join(["?"] * len(status_filter))
        clauses.append(f"p.status IN ({placeholders})")
        params.extend(status_filter)

    cleaned_rubric = rubric_query.strip().lower()
    if cleaned_rubric:
        clauses.append("LOWER(COALESCE(p.rubric, '')) LIKE ?")
        params.append(f"%{cleaned_rubric}%")

    cleaned_participant = participant_query.strip().lower()
    if cleaned_participant:
        clauses.append(
            """
            (
                LOWER(COALESCE(au.username, '')) LIKE ?
                OR LOWER(COALESCE(eu.username, '')) LIKE ?
                OR LOWER(COALESCE(pu.username, '')) LIKE ?
            )
            """
        )
        like_value = f"%{cleaned_participant}%"
        params.extend([like_value, like_value, like_value])

    if use_date_filter:
        clauses.append("DATE(p.created_at) >= DATE(?)")
        clauses.append("DATE(p.created_at) <= DATE(?)")
        params.append(date_from.isoformat())
        params.append(date_to.isoformat())

    where_sql = f"WHERE {' AND '.join(clauses)}" if clauses else ""

    rows = cursor.execute(
        f"""
        SELECT
            p.id,
            p.created_at,
            p.title,
            COALESCE(NULLIF(p.rubric, ''), '-') AS rubric,
            p.status,
            COALESCE(au.username, '-') AS author_name,
            COALESCE(eu.username, '-') AS executor_name,
            COALESCE(pu.username, '-') AS proofreader_name,
            COALESCE(p.is_archived, 0) AS is_archived
        FROM projects p
        LEFT JOIN users au ON au.id = COALESCE(p.author_user_id, p.author_id)
        LEFT JOIN users eu ON eu.id = p.executor_user_id
        LEFT JOIN users pu ON pu.id = p.proofreader_user_id
        {where_sql}
        ORDER BY p.created_at DESC, p.id DESC
        """,
        params,
    ).fetchall()

    result: list[dict[str, Any]] = []
    for row in rows:
        result.append(
            {
                "id": int(row[0]),
                "created_at": row[1],
                "title": row[2] or "",
                "rubric": row[3] or "-",
                "status": row[4] or "",
                "author_name": row[5] or "-",
                "executor_name": row[6] or "-",
                "proofreader_name": row[7] or "-",
                "is_archived": int(row[8]),
            }
        )
    return result


def render_main_table(project_rows: list[dict[str, Any]]) -> None:
    if not project_rows:
        st.info("По выбранным фильтрам пока нет проектов.")
        return

    header_cells = [
        "Дата создания",
        "Название",
        "Рубрика",
        "Статус",
        "Автор",
        "Исполнитель",
        "Корректор",
    ]

    rows_html: list[str] = []
    for row in project_rows:
        status_value = row["status"] or "draft"
        badge_class = f"nn-status-chip nn-status-{status_value}"
        rows_html.append(
            "<tr>"
            f"<td>{html.escape(format_date_only(row['created_at']))}</td>"
            f"<td><b>#{row['id']}</b> {html.escape(row['title'])}</td>"
            f"<td>{html.escape(row['rubric'])}</td>"
            f"<td><span class='{badge_class}'>{html.escape(status_label(status_value))}</span></td>"
            f"<td>{html.escape(row['author_name'])}</td>"
            f"<td>{html.escape(row['executor_name'])}</td>"
            f"<td>{html.escape(row['proofreader_name'])}</td>"
            "</tr>"
        )

    table_html = (
        """
        <style>
        .nn-main-table {
            width: 100%;
            border-collapse: collapse;
            font-size: 14px;
            background: #fff;
        }
        .nn-main-table th, .nn-main-table td {
            border: 1px solid #aaaaaa;
            padding: 7px 8px;
            text-align: left;
            vertical-align: top;
        }
        .nn-main-table th {
            background-color: #f7f7f7;
            font-weight: 700;
            font-size: 13px;
        }
        .nn-main-table tbody tr:nth-child(even) {
            background: #f5f5f5;
        }
        </style>
        """
        + "<table class='nn-main-table'>"
        + "<thead><tr>"
        + "".join(f"<th>{html.escape(cell)}</th>" for cell in header_cells)
        + "</tr></thead>"
        + "<tbody>"
        + "".join(rows_html)
        + "</tbody></table>"
    )

    st.markdown(table_html, unsafe_allow_html=True)


def fetch_projects_for_archive(
    cursor: sqlite3.Cursor,
    *,
    title_query: str,
    rubric_query: str,
    archived_by_query: str,
    use_date_filter: bool,
    date_from: date,
    date_to: date,
) -> list[dict[str, Any]]:
    clauses = ["(COALESCE(p.is_archived, 0) = 1 OR p.status = 'archived')"]
    params: list[Any] = []

    cleaned_title = title_query.strip().lower()
    if cleaned_title:
        clauses.append(
            "(LOWER(COALESCE(p.title, '')) LIKE ? OR LOWER(COALESCE(p.topic, '')) LIKE ?)"
        )
        title_like = f"%{cleaned_title}%"
        params.extend([title_like, title_like])

    cleaned_rubric = rubric_query.strip().lower()
    if cleaned_rubric:
        clauses.append("LOWER(COALESCE(p.rubric, '')) LIKE ?")
        params.append(f"%{cleaned_rubric}%")

    cleaned_archived_by = archived_by_query.strip().lower()
    if cleaned_archived_by:
        clauses.append("LOWER(COALESCE(ar.username, '')) LIKE ?")
        params.append(f"%{cleaned_archived_by}%")

    if use_date_filter:
        clauses.append("DATE(COALESCE(p.archived_at, p.created_at)) >= DATE(?)")
        clauses.append("DATE(COALESCE(p.archived_at, p.created_at)) <= DATE(?)")
        params.extend([date_from.isoformat(), date_to.isoformat()])

    where_sql = f"WHERE {' AND '.join(clauses)}"
    rows = cursor.execute(
        f"""
        SELECT
            p.id,
            p.title,
            COALESCE(NULLIF(p.rubric, ''), '-') AS rubric,
            p.created_at,
            p.archived_at,
            COALESCE(au.username, '-') AS author_name,
            COALESCE(ar.username, '-') AS archived_by_name
        FROM projects p
        LEFT JOIN users au ON au.id = COALESCE(p.author_user_id, p.author_id)
        LEFT JOIN users ar ON ar.id = p.archived_by
        {where_sql}
        ORDER BY COALESCE(p.archived_at, p.created_at) DESC, p.id DESC
        """,
        params,
    ).fetchall()

    result: list[dict[str, Any]] = []
    for row in rows:
        result.append(
            {
                "id": int(row[0]),
                "title": row[1] or "",
                "rubric": row[2] or "-",
                "created_at": row[3],
                "archived_at": row[4],
                "author_name": row[5] or "-",
                "archived_by_name": row[6] or "-",
            }
        )
    return result


def render_archive_table(project_rows: list[dict[str, Any]]) -> None:
    if not project_rows:
        st.info("Архив пуст или нет данных по текущим фильтрам.")
        return

    header_cells = [
        "Дата архивации",
        "Дата создания",
        "Название",
        "Рубрика",
        "Автор",
        "Кто архивировал",
    ]

    rows_html: list[str] = []
    for row in project_rows:
        rows_html.append(
            "<tr>"
            f"<td>{html.escape(format_datetime(row['archived_at']))}</td>"
            f"<td>{html.escape(format_date_only(row['created_at']))}</td>"
            f"<td><b>#{row['id']}</b> {html.escape(row['title'])}</td>"
            f"<td>{html.escape(row['rubric'])}</td>"
            f"<td>{html.escape(row['author_name'])}</td>"
            f"<td>{html.escape(row['archived_by_name'])}</td>"
            "</tr>"
        )

    table_html = (
        """
        <style>
        .nn-archive-table {
            width: 100%;
            border-collapse: collapse;
            font-size: 14px;
            background: #fff;
        }
        .nn-archive-table th, .nn-archive-table td {
            border: 1px solid #aaaaaa;
            padding: 7px 8px;
            text-align: left;
            vertical-align: top;
        }
        .nn-archive-table th {
            background-color: #f7f7f7;
            font-weight: 700;
        }
        .nn-archive-table tbody tr:nth-child(even) {
            background: #f5f5f5;
        }
        </style>
        """
        + "<table class='nn-archive-table'>"
        + "<thead><tr>"
        + "".join(f"<th>{html.escape(cell)}</th>" for cell in header_cells)
        + "</tr></thead>"
        + "<tbody>"
        + "".join(rows_html)
        + "</tbody></table>"
    )
    st.markdown(table_html, unsafe_allow_html=True)


def event_type_label(event_type: str) -> str:
    labels = {
        "project_created": "Проект создан",
        "project_cloned": "Проект скопирован",
        "status_changed": "Статус изменен",
        "project_archived": "Проект отправлен в архив",
        "project_restored": "Проект возвращен из архива",
        "file_uploaded": "Файл загружен",
        "file_deleted": "Файл удален",
    }
    return labels.get(event_type, event_type)


def fetch_project_event_history(
    cursor: sqlite3.Cursor,
    project_id: int,
    *,
    limit: int = 30,
) -> list[dict[str, Any]]:
    rows = cursor.execute(
        """
        SELECT
            pe.event_type,
            pe.old_value,
            pe.new_value,
            pe.created_at,
            COALESCE(u.username, '-') AS actor_name
        FROM project_events pe
        LEFT JOIN users u ON u.id = pe.actor_user_id
        WHERE pe.project_id = ?
        ORDER BY pe.id DESC
        LIMIT ?
        """,
        (project_id, limit),
    ).fetchall()

    return [
        {
            "event_type": row[0] or "",
            "old_value": row[1],
            "new_value": row[2],
            "created_at": row[3],
            "actor_name": row[4] or "-",
        }
        for row in rows
    ]


def restore_project_from_archive(
    conn: sqlite3.Connection,
    cursor: sqlite3.Cursor,
    project_id: int,
    actor_user_id: int,
) -> tuple[bool, str]:
    project_row = cursor.execute(
        "SELECT status, COALESCE(is_archived, 0) FROM projects WHERE id = ?",
        (project_id,),
    ).fetchone()
    if not project_row:
        return False, "Проект не найден."

    current_status, is_archived = project_row
    if int(is_archived) == 0 and current_status != "archived":
        return False, "Проект уже находится в рабочем списке."

    restore_status = "draft"
    previous_status_row = cursor.execute(
        """
        SELECT old_value
        FROM project_events
        WHERE project_id = ?
          AND event_type = 'status_changed'
          AND new_value = 'archived'
        ORDER BY id DESC
        LIMIT 1
        """,
        (project_id,),
    ).fetchone()
    if previous_status_row and previous_status_row[0] in STATUS_VALUES:
        restore_status = previous_status_row[0]
    if restore_status == "archived":
        restore_status = "draft"

    now = iso_now()
    cursor.execute(
        """
        UPDATE projects
        SET
            is_archived = 0,
            archived_at = NULL,
            archived_by = NULL,
            status = ?,
            status_changed_at = ?,
            status_changed_by = ?
        WHERE id = ?
        """,
        (restore_status, now, actor_user_id, project_id),
    )
    log_project_event(
        cursor,
        project_id=project_id,
        event_type="status_changed",
        actor_user_id=actor_user_id,
        old_value="archived",
        new_value=restore_status,
    )
    log_project_event(
        cursor,
        project_id=project_id,
        event_type="project_restored",
        actor_user_id=actor_user_id,
    )
    conn.commit()
    return True, restore_status


init_db()
inject_global_ui_css()

if "user" not in st.session_state:
    auth_token = read_auth_token_from_query()
    if auth_token:
        restored_user_id = verify_session_token(auth_token)
        if restored_user_id is not None:
            restore_conn = get_conn()
            restore_cursor = restore_conn.cursor()
            restored_user = fetch_user_by_id(restore_cursor, restored_user_id)
            restore_conn.close()
            if restored_user:
                st.session_state.user = restored_user
                # Keep token in URL and refresh its TTL on successful restore.
                write_auth_token_to_query(create_session_token(restored_user["id"]))
            else:
                clear_auth_token_from_query()
        else:
            clear_auth_token_from_query()

if "user" not in st.session_state:
    render_screen_header("LOGIN")
    st.subheader("Вход")

    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        auth_user = authenticate(username, password)
        if auth_user:
            st.session_state.user = auth_user
            write_auth_token_to_query(create_session_token(auth_user["id"]))
            st.rerun()
        else:
            st.error("Invalid credentials")
    st.stop()

user = st.session_state.user

with st.sidebar:
    st.subheader("Профиль")
    st.write(f"Пользователь: `{user['username']}`")
    st.write(f"Роль: `{user['role']}`")
    section_options = ["MAIN", "EDITOR", "ARHIV"]
    section_value = st.radio(
        "Раздел",
        section_options,
        index=option_index(section_options, st.session_state.get("current_section", "MAIN")),
    )
    st.session_state.current_section = section_value
    if st.button("Выйти из профиля"):
        del st.session_state.user
        clear_auth_token_from_query()
        st.rerun()

if st.session_state.current_section == "ARHIV":
    render_screen_header("ARHIV", "Архивные проекты: фильтры, возврат в работу и история событий.")
    render_archive_action_blocks()

    conn = get_conn()
    cursor = conn.cursor()

    can_restore = user["role"] in ("admin", "editor")
    archive_filter_col_1, archive_filter_col_2, archive_filter_col_3 = st.columns(3)
    with archive_filter_col_1:
        archive_title_query = st.text_input("Название/тема содержит", value="")
    with archive_filter_col_2:
        archive_rubric_query = st.text_input("Рубрика содержит", value="")
    with archive_filter_col_3:
        archive_by_query = st.text_input("Кто архивировал (логин)", value="")

    use_archive_date_filter = st.checkbox("Фильтр по дате архивации", value=False)
    archive_date_from, archive_date_to = date.today(), date.today()
    if use_archive_date_filter:
        archive_date_col_1, archive_date_col_2 = st.columns(2)
        with archive_date_col_1:
            archive_date_from = st.date_input("Дата от", value=date.today(), key="arch_date_from")
        with archive_date_col_2:
            archive_date_to = st.date_input("Дата до", value=date.today(), key="arch_date_to")
        if archive_date_from > archive_date_to:
            st.error("Дата 'от' не может быть больше даты 'до'.")
            conn.close()
            st.stop()

    archived_projects = fetch_projects_for_archive(
        cursor,
        title_query=archive_title_query,
        rubric_query=archive_rubric_query,
        archived_by_query=archive_by_query,
        use_date_filter=use_archive_date_filter,
        date_from=archive_date_from,
        date_to=archive_date_to,
    )

    st.markdown("### Таблица архива")
    render_archive_table(archived_projects)

    if not archived_projects:
        conn.close()
        st.stop()

    archived_project_ids = [row["id"] for row in archived_projects]
    archived_project_by_id = {row["id"]: row for row in archived_projects}

    selected_archived_project_id = st.session_state.get(
        "selected_archived_project_id",
        archived_project_ids[0],
    )
    if selected_archived_project_id not in archived_project_ids:
        selected_archived_project_id = archived_project_ids[0]

    selected_archived_project_id = st.selectbox(
        "Выбрать архивный проект",
        archived_project_ids,
        index=archived_project_ids.index(selected_archived_project_id),
        format_func=lambda project_id: (
            f"#{project_id} | {archived_project_by_id[project_id]['title']} "
            f"| архив: {format_datetime(archived_project_by_id[project_id]['archived_at'])}"
        ),
    )
    st.session_state.selected_archived_project_id = selected_archived_project_id

    if can_restore:
        if st.button("↩️ Вернуть выбранный проект в работу", use_container_width=True):
            st.session_state.restore_confirm_project_id = selected_archived_project_id
    else:
        st.info("Возвращать проект из архива могут только роли admin/editor.")

    restore_confirm_id = st.session_state.get("restore_confirm_project_id")
    if restore_confirm_id == selected_archived_project_id and can_restore:
        st.warning(
            f"Подтвердите возврат проекта #{selected_archived_project_id} из архива в MAIN."
        )
        restore_col_1, restore_col_2 = st.columns(2)
        with restore_col_1:
            if st.button("Да, вернуть в работу", use_container_width=True):
                restored, restore_result = restore_project_from_archive(
                    conn,
                    cursor,
                    selected_archived_project_id,
                    user["id"],
                )
                st.session_state.restore_confirm_project_id = None
                if restored:
                    st.success(
                        f"Проект #{selected_archived_project_id} возвращен в работу "
                        f"(статус: {status_label(restore_result)})."
                    )
                else:
                    st.error(restore_result)
                st.rerun()
        with restore_col_2:
            if st.button("Отмена", use_container_width=True):
                st.session_state.restore_confirm_project_id = None
                st.rerun()

    st.markdown("### История событий выбранного проекта")
    history_rows = fetch_project_event_history(cursor, selected_archived_project_id)
    if not history_rows:
        st.info("По этому проекту пока нет событий истории.")
    else:
        for history in history_rows:
            event_label = event_type_label(history["event_type"])
            old_value = history["old_value"] or "-"
            new_value = history["new_value"] or "-"
            st.write(
                f"- [{format_datetime(history['created_at'])}] "
                f"{history['actor_name']} | {event_label} | {old_value} -> {new_value}"
            )

    conn.close()
    st.stop()

conn = get_conn()
cursor = conn.cursor()

if st.session_state.current_section == "MAIN":
    render_screen_header("MAIN", "Рабочий список сюжетов: создание, фильтры, архивирование и переход в text-editor.")

    can_create = user["role"] in ("author", "admin", "editor")
    can_archive = user["role"] in ("admin", "editor")

    st.markdown("### Фильтры MAIN")
    filter_col_1, filter_col_2, filter_col_3, filter_col_4 = st.columns(4)

    with filter_col_1:
        include_archived = st.checkbox("Показывать архив", value=False)
    with filter_col_2:
        status_filter = st.multiselect(
            "Статусы",
            STATUS_VALUES,
            default=[],
            format_func=status_label,
        )
    with filter_col_3:
        rubric_query = st.text_input("Рубрика содержит", value="")
    with filter_col_4:
        participant_query = st.text_input("Участник содержит", value="")

    use_date_filter = st.checkbox("Фильтр по дате создания", value=False)
    date_from, date_to = date.today(), date.today()
    if use_date_filter:
        date_col_1, date_col_2 = st.columns(2)
        with date_col_1:
            date_from = st.date_input("Дата от", value=date.today())
        with date_col_2:
            date_to = st.date_input("Дата до", value=date.today())
        if date_from > date_to:
            st.error("Дата 'от' не может быть больше даты 'до'.")
            conn.close()
            st.stop()

    projects = fetch_projects_for_main(
        cursor,
        include_archived=include_archived,
        status_filter=status_filter,
        rubric_query=rubric_query,
        participant_query=participant_query,
        use_date_filter=use_date_filter,
        date_from=date_from,
        date_to=date_to,
    )

    st.markdown("### Таблица проектов")
    render_main_table(projects)

    if not projects:
        conn.close()
        st.stop()

    project_ids = [row["id"] for row in projects]
    project_by_id = {row["id"]: row for row in projects}

    selected_project_id = st.session_state.get("selected_project_id", project_ids[0])
    if selected_project_id not in project_ids:
        selected_project_id = project_ids[0]

    selected_project_id = st.selectbox(
        "Выбранный проект",
        options=project_ids,
        index=project_ids.index(selected_project_id),
        format_func=lambda project_id: (
            f"#{project_id} | {project_by_id[project_id]['title']} "
            f"| {format_date_only(project_by_id[project_id]['created_at'])}"
        ),
    )
    st.session_state.selected_project_id = selected_project_id

    quick_export_payload = fetch_project_export_payload(cursor, selected_project_id)
    quick_docx_bytes, _quick_docx_error = generate_docx_bytes(quick_export_payload)
    quick_pdf_bytes, _quick_pdf_error = generate_pdf_bytes(quick_export_payload)

    st.markdown("### Действия MAIN (как в вашем табличном шаблоне)")
    action_col_1, action_col_2, action_col_3, action_col_4 = st.columns(4)
    action_col_5, action_col_6, action_col_7, action_col_8 = st.columns(4)

    with action_col_1:
        if can_create and st.button("Создать новый пустой", use_container_width=True):
            new_project_id = create_empty_project(conn, cursor, user["id"])
            st.session_state.selected_project_id = new_project_id
            st.success(f"Создан новый проект #{new_project_id}.")
            st.rerun()
    with action_col_2:
        if can_create and st.button("Создать на основе последнего", use_container_width=True):
            last_row = cursor.execute(
                """
                SELECT id
                FROM projects
                WHERE COALESCE(is_archived, 0) = 0
                ORDER BY created_at DESC, id DESC
                LIMIT 1
                """
            ).fetchone()
            if not last_row:
                st.warning("Нет проекта, который можно клонировать.")
            else:
                new_project_id = clone_project(conn, cursor, int(last_row[0]), user["id"])
                if new_project_id is None:
                    st.error("Не удалось создать копию проекта.")
                else:
                    st.session_state.selected_project_id = new_project_id
                    st.success(f"Создана копия проекта #{new_project_id}.")
                    st.rerun()
    with action_col_3:
        if can_create and st.button("Создать на основе выбранного", use_container_width=True):
            new_project_id = clone_project(conn, cursor, selected_project_id, user["id"])
            if new_project_id is None:
                st.error("Не удалось создать копию выбранного проекта.")
            else:
                st.session_state.selected_project_id = new_project_id
                st.success(f"Создана копия проекта #{new_project_id}.")
                st.rerun()
    with action_col_4:
        if quick_docx_bytes is not None:
            st.download_button(
                "Export Word",
                data=quick_docx_bytes,
                file_name=f"newscast_project_{selected_project_id}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.button("Export Word", disabled=True, use_container_width=True)
    with action_col_5:
        if quick_pdf_bytes is not None:
            st.download_button(
                "Export PDF",
                data=quick_pdf_bytes,
                file_name=f"newscast_project_{selected_project_id}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )
        else:
            st.button("Export PDF", disabled=True, use_container_width=True)
    with action_col_6:
        if st.button("Открыть ARCHIV", use_container_width=True):
            st.session_state.current_section = "ARHIV"
            st.rerun()
    with action_col_7:
        if can_archive and st.button("Сдать в архив", use_container_width=True):
            st.session_state.archive_confirm_project_id = selected_project_id
    with action_col_8:
        if st.button("Открыть text-editor", use_container_width=True):
            st.session_state.current_section = "EDITOR"
            st.rerun()

    confirm_project_id = st.session_state.get("archive_confirm_project_id")
    if confirm_project_id == selected_project_id:
        st.warning(
            f"Подтвердите архивирование проекта #{selected_project_id}. "
            "После этого он уйдет в ARHIV."
        )
        confirm_col_1, confirm_col_2 = st.columns(2)
        with confirm_col_1:
            if st.button("Да, архивировать", use_container_width=True):
                archived = archive_project(conn, cursor, selected_project_id, user["id"])
                st.session_state.archive_confirm_project_id = None
                if archived:
                    st.success(f"Проект #{selected_project_id} отправлен в архив.")
                else:
                    st.info("Проект уже в архиве или не найден.")
                st.rerun()
        with confirm_col_2:
            if st.button("Отмена", use_container_width=True):
                st.session_state.archive_confirm_project_id = None
                st.rerun()

    conn.close()
    st.stop()

if st.session_state.current_section != "EDITOR":
    conn.close()
    st.stop()

render_screen_header("text-editor", "Отдельный экран редактирования выбранного проекта.")
editor_nav_col_1, editor_nav_col_2 = st.columns(2)
with editor_nav_col_1:
    if st.button("← Вернуться в MAIN", use_container_width=True):
        st.session_state.current_section = "MAIN"
        st.rerun()
with editor_nav_col_2:
    if st.button("Перейти в ARHIV", use_container_width=True):
        st.session_state.current_section = "ARHIV"
        st.rerun()

editor_projects = fetch_projects_for_main(
    cursor,
    include_archived=False,
    status_filter=[],
    rubric_query="",
    participant_query="",
    use_date_filter=False,
    date_from=date.today(),
    date_to=date.today(),
)
if not editor_projects:
    st.info("Нет рабочих проектов для редактирования. Создайте проект в MAIN.")
    conn.close()
    st.stop()

editor_project_ids = [row["id"] for row in editor_projects]
editor_project_by_id = {row["id"]: row for row in editor_projects}
selected_project_id = st.session_state.get("selected_project_id", editor_project_ids[0])
if selected_project_id not in editor_project_ids:
    selected_project_id = editor_project_ids[0]

selected_project_id = st.selectbox(
    "Проект для редактирования",
    editor_project_ids,
    index=editor_project_ids.index(selected_project_id),
    format_func=lambda project_id: (
        f"#{project_id} | {editor_project_by_id[project_id]['title']} "
        f"| {format_date_only(editor_project_by_id[project_id]['created_at'])}"
    ),
)
st.session_state.selected_project_id = selected_project_id

project_data = cursor.execute(
    """
    SELECT
        title,
        status,
        COALESCE(NULLIF(rubric, ''), '') AS rubric,
        COALESCE(NULLIF(planned_duration, ''), '') AS planned_duration,
        COALESCE(author_user_id, author_id) AS author_user_id,
        executor_user_id,
        proofreader_user_id,
        COALESCE(file_root, '') AS file_root
    FROM projects
    WHERE id = ?
    """,
    (selected_project_id,),
).fetchone()

if not project_data:
    conn.close()
    st.error("Выбранный проект не найден.")
    st.stop()

(
    title,
    status,
    rubric,
    planned_duration,
    author_user_id,
    executor_user_id,
    proofreader_user_id,
    file_root,
) = project_data

project_storage_dir = resolve_project_storage_dir(selected_project_id, file_root)

users = fetch_users(cursor)
users_map = {user_row["id"]: user_row for user_row in users}
assignee_options: list[int | None] = [None] + [user_row["id"] for user_row in users]

editor_editable = can_edit(user, status)
meta_editable = user["role"] in PROJECT_META_EDIT_ROLES
assignment_editable = user["role"] in PROJECT_ASSIGN_EDIT_ROLES
status_editable = user["role"] in ("admin", "editor", "proofreader")

st.caption("Верхний блок и кнопки выровнены под ваш табличный паттерн; рабочие формы находятся ниже.")

export_payload = fetch_project_export_payload(cursor, selected_project_id)
docx_bytes, docx_error = generate_docx_bytes(export_payload)
pdf_bytes, pdf_error = generate_pdf_bytes(export_payload)

control_col_1, control_col_2, control_col_3, control_col_4 = st.columns(4)
with control_col_1:
    titling_mode = st.checkbox(
        "Режим титрования",
        value=st.session_state.get("titling_mode", False),
        help="Скрывает лишние колонки и оставляет только то, что нужно для титров.",
    )
    st.session_state.titling_mode = titling_mode
with control_col_2:
    print_preview = st.checkbox(
        "Печатная версия",
        value=st.session_state.get("print_preview_mode", False),
        help="Показывает подготовленную версию для печати/сохранения в PDF из браузера.",
    )
    st.session_state.print_preview_mode = print_preview
with control_col_3:
    if docx_bytes is not None:
        st.download_button(
            "📄 Export Word",
            data=docx_bytes,
            file_name=f"newscast_project_{selected_project_id}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    else:
        st.button("📄 Export Word", disabled=True, use_container_width=True)
with control_col_4:
    if pdf_bytes is not None:
        st.download_button(
            "📕 Export PDF",
            data=pdf_bytes,
            file_name=f"newscast_project_{selected_project_id}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    else:
        st.button("📕 Export PDF", disabled=True, use_container_width=True)

if docx_error:
    st.info(docx_error)
if pdf_error:
    st.info(pdf_error)

st.caption(f"Локальное хранилище проекта: `{project_storage_dir}`")

if st.session_state.get("print_preview_mode"):
    st.markdown(
        """
        <style>
        @media print {
            @page {
                size: auto;
                margin: 10mm;
            }

            html, body {
                margin: 0 !important;
                padding: 0 !important;
                height: auto !important;
                overflow: visible !important;
                background: white !important;
            }

            /* Remove all Streamlit blocks from print flow by default. */
            [data-testid="stAppViewContainer"] .element-container {
                display: none !important;
            }

            /* Keep only block that contains our simplified print markup. */
            [data-testid="stAppViewContainer"] .element-container:has(.nn-print-target) {
                display: block !important;
            }

            [data-testid="stAppViewContainer"] .main .block-container {
                padding-top: 0 !important;
                padding-bottom: 0 !important;
            }

            .nn-print-target {
                display: block !important;
                position: static !important;
                width: 100% !important;
                margin: 0 !important;
                padding: 0 !important;
                page-break-after: auto;
            }

            .nn-print-target .nn-print-table tr {
                break-inside: avoid;
                page-break-inside: avoid;
            }

            [data-testid="stSidebar"],
            header,
            footer {
                display: none !important;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.markdown("#### Печатная версия")
    st.caption("Нажмите Cmd+P (или Ctrl+P) — печататься будет только блок ниже.")
    st.markdown(
        f"<div class='nn-print-target'>{build_print_html(export_payload)}</div>",
        unsafe_allow_html=True,
    )

with st.form("project_meta_form"):
    st.markdown("#### Метаданные проекта")
    meta_col_1, meta_col_2, meta_col_3 = st.columns(3)
    with meta_col_1:
        title_input = st.text_input(
            "Название",
            value=title or "",
            disabled=not meta_editable,
        )
        rubric_input = st.text_input(
            "Рубрика",
            value=rubric or "",
            disabled=not meta_editable,
        )
    with meta_col_2:
        duration_input = st.text_input(
            "Хронометраж",
            value=planned_duration or "",
            disabled=not meta_editable,
            help="Пример: 02:30 или 00:45",
        )
        file_root_input = st.text_input(
            "Путь к файлам проекта",
            value=file_root or "",
            disabled=not meta_editable,
            help="Локальный путь на сервере. Пока без загрузки файлов.",
        )
    with meta_col_3:
        status_input = st.selectbox(
            "Статус",
            STATUS_VALUES,
            index=option_index(STATUS_VALUES, status),
            format_func=status_label,
            disabled=not status_editable,
        )

    assign_col_1, assign_col_2, assign_col_3 = st.columns(3)
    with assign_col_1:
        author_input = st.selectbox(
            "Автор",
            assignee_options,
            index=option_index(assignee_options, author_user_id),
            format_func=lambda uid: "— не назначен —" if uid is None else user_label(users_map, uid),
            disabled=not assignment_editable,
        )
    with assign_col_2:
        executor_input = st.selectbox(
            "Исполнитель",
            assignee_options,
            index=option_index(assignee_options, executor_user_id),
            format_func=lambda uid: "— не назначен —" if uid is None else user_label(users_map, uid),
            disabled=not assignment_editable,
        )
    with assign_col_3:
        proofreader_input = st.selectbox(
            "Корректор",
            assignee_options,
            index=option_index(assignee_options, proofreader_user_id),
            format_func=lambda uid: "— не назначен —" if uid is None else user_label(users_map, uid),
            disabled=not assignment_editable,
        )

    save_meta = st.form_submit_button("Сохранить метаданные")
    if save_meta:
        changes_applied = False

        if meta_editable:
            clean_title = (title_input or "").strip() or "Новый проект"
            clean_rubric = (rubric_input or "").strip()
            clean_duration = (duration_input or "").strip()
            clean_file_root = (file_root_input or "").strip()
            cursor.execute(
                """
                UPDATE projects
                SET
                    title = ?,
                    rubric = ?,
                    planned_duration = ?,
                    file_root = ?
                WHERE id = ?
                """,
                (
                    clean_title,
                    clean_rubric,
                    clean_duration,
                    clean_file_root,
                    selected_project_id,
                ),
            )
            changes_applied = True

        if assignment_editable:
            cursor.execute(
                """
                UPDATE projects
                SET
                    author_user_id = ?,
                    author_id = ?,
                    executor_user_id = ?,
                    proofreader_user_id = ?
                WHERE id = ?
                """,
                (
                    author_input,
                    author_input,
                    executor_input,
                    proofreader_input,
                    selected_project_id,
                ),
            )
            changes_applied = True

        if status_editable and status_input != status:
            cursor.execute(
                """
                UPDATE projects
                SET
                    status = ?,
                    status_changed_at = ?,
                    status_changed_by = ?
                WHERE id = ?
                """,
                (status_input, iso_now(), user["id"], selected_project_id),
            )
            log_project_event(
                cursor,
                project_id=selected_project_id,
                event_type="status_changed",
                actor_user_id=user["id"],
                old_value=status,
                new_value=status_input,
            )
            changes_applied = True

        if changes_applied:
            conn.commit()
            st.success("Метаданные проекта обновлены.")
            st.rerun()
        else:
            st.info("Нет прав для изменения метаданных этого проекта.")

elements = cursor.execute(
    """
    SELECT
        id,
        order_index,
        COALESCE(text, '') AS text,
        COALESCE(NULLIF(block_type, ''), 'zk') AS block_type,
        COALESCE(speaker_text, '') AS speaker_text,
        COALESCE(file_name, '') AS file_name,
        COALESCE(tc_in, '') AS tc_in,
        COALESCE(tc_out, '') AS tc_out,
        COALESCE(additional_comment, '') AS additional_comment
    FROM script_elements
    WHERE project_id = ?
    ORDER BY order_index
    """,
    (selected_project_id,),
).fetchall()

if pd is None:
    st.error("Для табличного редактора не найден pandas. Установите зависимость и перезапустите сервис.")
    conn.close()
    st.stop()

editor_df = build_editor_dataframe(elements)
if editor_df is None:
    st.error("Не удалось подготовить данные для таблицы.")
    conn.close()
    st.stop()

st.markdown("### Редактируемая таблица сценария")
st.caption(
    "Редактируйте ячейки напрямую. "
    "Для добавления строки используйте кнопку `Добавить строку` (номер ставится автоматически). "
    "Для удаления выделите строки кликом в этой таблице и нажмите `Удалить выбранные`. "
    "После изменений нажмите `Сохранить таблицу`."
)
st.caption("Для блока СНХ в колонке `Титр` укажите две строки: 1) ФИО, 2) должность.")

column_config: dict[str, Any] = {
    "ID": st.column_config.NumberColumn("ID", disabled=True, width="small"),
    "№": st.column_config.NumberColumn("№", disabled=True, width="small"),
    "Блок": st.column_config.SelectboxColumn("Блок", options=BLOCK_TYPE_LABELS, required=True, width="small"),
    "Текст": st.column_config.TextColumn("Текст", width="large"),
    "Титр": st.column_config.TextColumn("Титр", width="medium"),
    "Имя файла": st.column_config.TextColumn("Имя файла", width="medium"),
    "TC IN": st.column_config.TextColumn("TC IN", width="small"),
    "TC OUT": st.column_config.TextColumn("TC OUT", width="small"),
    "Другой коммент": st.column_config.TextColumn("Другой коммент", width="medium"),
}

# Hide technical DB key from UI, but keep it in dataframe for stable updates.
column_config["ID"] = None
if titling_mode:
    column_config["Имя файла"] = None
    column_config["TC IN"] = None
    column_config["TC OUT"] = None
    column_config["Другой коммент"] = None

table_key = f"editor_grid_v5_{selected_project_id}_{'titl' if titling_mode else 'full'}"
table_rows_current: list[dict[str, Any]] = []
selected_row_positions: list[int] = []

if AgGrid is None or GridOptionsBuilder is None or GridUpdateMode is None or DataReturnMode is None:
    st.warning(
        "Пакет `streamlit-aggrid` не установлен. "
        "Пока включен базовый режим без выбора строк кликом."
    )
    edited_df = st.data_editor(
        editor_df,
        key=table_key,
        hide_index=True,
        use_container_width=True,
        num_rows="fixed",
        disabled=not editor_editable,
        column_config=column_config,
        height=520,
    )
    table_rows_current = edited_df.to_dict(orient="records")
else:
    grid_df = editor_df.copy()
    grid_builder = GridOptionsBuilder.from_dataframe(grid_df)
    grid_builder.configure_default_column(
        editable=editor_editable,
        resizable=True,
        sortable=False,
        filter=False,
        wrapText=True,
        autoHeight=True,
    )
    grid_builder.configure_column("ID", hide=True, editable=False)
    grid_builder.configure_column("№", editable=False)
    grid_builder.configure_column(
        "Блок",
        cellEditor="agSelectCellEditor",
        cellEditorParams={"values": BLOCK_TYPE_LABELS},
    )
    if titling_mode:
        grid_builder.configure_column("Имя файла", hide=True, editable=False)
        grid_builder.configure_column("TC IN", hide=True, editable=False)
        grid_builder.configure_column("TC OUT", hide=True, editable=False)
        grid_builder.configure_column("Другой коммент", hide=True, editable=False)
    grid_builder.configure_selection(
        selection_mode="multiple",
        use_checkbox=False,
        rowMultiSelectWithClick=True,
    )
    grid_builder.configure_grid_options(
        rowSelection="multiple",
        rowMultiSelectWithClick=True,
        suppressRowClickSelection=False,
        suppressCellFocus=False,
    )
    grid_response = AgGrid(
        grid_df,
        gridOptions=grid_builder.build(),
        data_return_mode=DataReturnMode.AS_INPUT,
        update_mode=GridUpdateMode.MODEL_CHANGED | GridUpdateMode.SELECTION_CHANGED,
        fit_columns_on_grid_load=True,
        allow_unsafe_jscode=False,
        height=520,
        theme="streamlit",
        key=table_key,
    )

    raw_grid_data = grid_response.get("data", grid_df)
    if pd is not None and isinstance(raw_grid_data, pd.DataFrame):
        table_rows_current = raw_grid_data.to_dict(orient="records")
    elif isinstance(raw_grid_data, list):
        table_rows_current = [dict(raw_row) for raw_row in raw_grid_data if isinstance(raw_row, dict)]
    else:
        table_rows_current = grid_df.to_dict(orient="records")

    for table_row in table_rows_current:
        if isinstance(table_row, dict):
            table_row.pop("_selectedRowNodeInfo", None)

    selected_rows_raw = grid_response.get("selected_rows", [])
    if pd is not None and isinstance(selected_rows_raw, pd.DataFrame):
        selected_rows = selected_rows_raw.to_dict(orient="records")
    elif isinstance(selected_rows_raw, list):
        selected_rows = selected_rows_raw
    else:
        selected_rows = []

    selected_row_positions_set: set[int] = set()
    selected_row_ids: set[int] = set()
    for selected_row in selected_rows:
        if not isinstance(selected_row, dict):
            continue
        selected_row_id = _coerce_int_or_none(selected_row.get("ID"))
        if selected_row_id is not None:
            selected_row_ids.add(selected_row_id)
        node_info = selected_row.get("_selectedRowNodeInfo")
        if isinstance(node_info, dict):
            node_row_index = node_info.get("nodeRowIndex")
            if isinstance(node_row_index, int):
                selected_row_positions_set.add(node_row_index)

    if not selected_row_positions_set and selected_row_ids:
        for row_position, table_row in enumerate(table_rows_current):
            table_row_id = _coerce_int_or_none(table_row.get("ID"))
            if table_row_id is not None and table_row_id in selected_row_ids:
                selected_row_positions_set.add(row_position)

    selected_row_positions = sorted(
        row_position
        for row_position in selected_row_positions_set
        if 0 <= row_position < len(table_rows_current)
    )

st.caption("Выделение для удаления: клик по строке; множественный выбор: Ctrl/Cmd + клик.")
selected_row_numbers = [
    str(_coerce_int_or_none(table_rows_current[row_position].get("№")) or (row_position + 1))
    for row_position in selected_row_positions
]
if selected_row_numbers:
    st.caption(f"Выбрано строк: {len(selected_row_numbers)} (№ {', '.join(selected_row_numbers)})")
else:
    st.caption("Выбрано строк: 0")

save_col_1, save_col_2, save_col_3, save_col_4 = st.columns([1.4, 1.4, 1.6, 2.0])
with save_col_1:
    save_table_clicked = st.button(
        "Сохранить таблицу",
        type="primary",
        use_container_width=True,
        disabled=not editor_editable,
    )
with save_col_2:
    add_row_clicked = st.button(
        "Добавить строку",
        use_container_width=True,
        disabled=not editor_editable,
    )
with save_col_3:
    delete_rows_clicked = st.button(
        "Удалить выбранные",
        use_container_width=True,
        disabled=not editor_editable or not selected_row_positions,
    )
with save_col_4:
    if not editor_editable:
        st.info("Текущий статус/роль не позволяет редактировать таблицу сценария.")
    else:
        st.caption("После сохранения порядок строк фиксируется по текущему порядку в таблице.")

if add_row_clicked:
    next_index_row = cursor.execute(
        "SELECT COALESCE(MAX(order_index), 0) + 1 FROM script_elements WHERE project_id = ?",
        (selected_project_id,),
    ).fetchone()
    next_index = int(next_index_row[0] or 1)
    cursor.execute(
        """
        INSERT INTO script_elements (
            project_id, order_index, text, element_type, block_type,
            speaker_text, file_name, tc_in, tc_out, additional_comment
        ) VALUES (?, ?, '', 'zk', 'zk', '', '', '', '', '')
        """,
        (
            selected_project_id,
            next_index,
        ),
    )
    conn.commit()
    st.success(f"Добавлена новая строка №{next_index}.")
    st.rerun()

if delete_rows_clicked:
    if not selected_row_positions:
        st.warning("Сначала выделите хотя бы одну строку кликом в таблице.")
    else:
        selected_row_positions_set = set(selected_row_positions)
        remaining_rows = [
            dict(row)
            for row_position, row in enumerate(table_rows_current)
            if row_position not in selected_row_positions_set
        ]

        if titling_mode:
            full_row_by_id = {
                _coerce_int_or_none(row.get("ID")): row
                for row in editor_df.to_dict(orient="records")
            }
            for table_row in remaining_rows:
                row_id = _coerce_int_or_none(table_row.get("ID"))
                full_row = full_row_by_id.get(row_id, {})
                table_row["Имя файла"] = full_row.get("Имя файла", "")
                table_row["TC IN"] = full_row.get("TC IN", "")
                table_row["TC OUT"] = full_row.get("TC OUT", "")
                table_row["Другой коммент"] = full_row.get("Другой коммент", "")

        normalized_rows, table_errors = normalize_editor_rows(remaining_rows)
        if table_errors:
            for table_error in table_errors:
                st.error(table_error)
        else:
            updated_count, inserted_count, removed_count = save_editor_rows(
                conn,
                cursor,
                project_id=selected_project_id,
                rows=normalized_rows,
            )
            st.success(
                "Удаление выполнено: "
                f"удалено {removed_count}, обновлено {updated_count}, "
                f"добавлено {inserted_count}."
            )
            st.rerun()

if save_table_clicked:
    table_rows = [dict(row) for row in table_rows_current]
    if titling_mode:
        full_row_by_id = {
            _coerce_int_or_none(row.get("ID")): row
            for row in editor_df.to_dict(orient="records")
        }
        for table_row in table_rows:
            row_id = _coerce_int_or_none(table_row.get("ID"))
            full_row = full_row_by_id.get(row_id, {})
            table_row["Имя файла"] = full_row.get("Имя файла", "")
            table_row["TC IN"] = full_row.get("TC IN", "")
            table_row["TC OUT"] = full_row.get("TC OUT", "")
            table_row["Другой коммент"] = full_row.get("Другой коммент", "")

    normalized_rows, table_errors = normalize_editor_rows(table_rows)
    if table_errors:
        for table_error in table_errors:
            st.error(table_error)
    else:
        updated_count, inserted_count, removed_count = save_editor_rows(
            conn,
            cursor,
            project_id=selected_project_id,
            rows=normalized_rows,
        )
        st.success(
            "Таблица сохранена: "
            f"обновлено {updated_count}, добавлено {inserted_count}, "
            f"удалено {removed_count} строк."
        )
        st.rerun()

st.markdown("### Комментарии проекта")
if not comments_have_project_scope(cursor):
    st.warning(
        "Для проектных комментариев нужно применить новую миграцию БД. "
        "Запустите: `python3 scripts/migrate_db.py` и перезапустите приложение."
    )
else:
    project_comment_rows = cursor.execute(
        """
        SELECT c.text, c.created_at, COALESCE(u.username, 'unknown')
        FROM comments c
        LEFT JOIN users u ON u.id = c.user_id
        WHERE c.project_id = ? AND c.element_id IS NULL
        ORDER BY c.id DESC
        """,
        (selected_project_id,),
    ).fetchall()
    if project_comment_rows:
        for comment_text, comment_created_at, comment_author in project_comment_rows:
            created_label = format_datetime(comment_created_at)
            st.write(f"- [{created_label}] {comment_author}: {comment_text}")
    else:
        st.caption("Комментариев проекта пока нет.")

    if editor_editable:
        with st.form("project_comment_form"):
            project_comment_input = st.text_area(
                "Добавить комментарий к проекту",
                key="new_comment_for_project",
                height=90,
            )
            save_project_comment = st.form_submit_button("Сохранить комментарий проекта")
            if save_project_comment and (project_comment_input or "").strip():
                cursor.execute(
                    """
                    INSERT INTO comments (project_id, element_id, user_id, text, created_at)
                    VALUES (?, NULL, ?, ?, ?)
                    """,
                    (selected_project_id, user["id"], project_comment_input.strip(), iso_now()),
                )
                conn.commit()
                st.success("Комментарий проекта добавлен.")
                st.rerun()

st.markdown("### Файлы проекта")
project_files = fetch_project_files(cursor, selected_project_id)
if project_files:
    for file_row in project_files:
        file_cols = st.columns([3.2, 1.0, 1.4, 1.4, 1.0])
        file_cols[0].write(file_row["original_name"])
        file_cols[1].write(f"{file_row['file_size'] // 1024} KB")
        file_cols[2].write(file_row["uploaded_by_name"])
        file_cols[3].write(format_datetime(file_row["uploaded_at"]))
        if editor_editable and file_cols[4].button("Удалить", key=f"delete_project_file_{file_row['id']}"):
            removed, remove_message = remove_project_file(
                conn,
                cursor,
                project_id=selected_project_id,
                file_id=file_row["id"],
                actor_user_id=user["id"],
            )
            if removed:
                st.success(remove_message)
            else:
                st.error(remove_message)
            st.rerun()

        file_path = Path(file_row["storage_path"])
        file_status = "найден" if file_path.exists() else "отсутствует"
        st.caption(f"Путь: `{file_path}` | Статус на диске: {file_status}")
else:
    st.caption("Файлов проекта пока нет.")

if editor_editable:
    upload_project_file = st.file_uploader(
        "Загрузить файл к проекту",
        key=f"upload_file_project_{selected_project_id}",
        help=(
            f"Допустимые расширения: {', '.join(sorted(ALLOWED_UPLOAD_EXTENSIONS))}. "
            f"Лимит: {MAX_UPLOAD_SIZE_MB} MB."
        ),
    )
    if st.button("Сохранить файл проекта"):
        ok, validation_message = validate_uploaded_file(upload_project_file)
        if not ok:
            st.error(validation_message)
        else:
            try:
                destination_path, original_name, file_size = save_uploaded_file_to_disk(
                    upload_project_file,
                    project_storage_dir / "project_files",
                )
            except OSError as exc:
                st.error(f"Не удалось сохранить файл на диск: {exc}")
            else:
                add_project_file_record(
                    conn,
                    cursor,
                    project_id=selected_project_id,
                    element_id=None,
                    original_name=original_name,
                    storage_path=str(destination_path),
                    mime_type=getattr(upload_project_file, "type", "") or "",
                    file_size=file_size,
                    uploaded_by=user["id"],
                )
                st.success(f"Файл проекта `{original_name}` загружен.")
                st.rerun()

conn.close()
